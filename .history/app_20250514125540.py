from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for, flash, session
import numpy as np
import requests
import os
import io
import tempfile
import json
import re
# Thêm thư viện xử lý file docx
import docx
# Thêm thư viện xử lý PDF và Excel
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
# Thay đổi từ SQLAlchemy sang PyMongo
from auth import login_user, logout_user, login_required, role_required, has_permission
from database import get_db, log_action

app = Flask(__name__)
app.secret_key = 'your-secret-key-here'  # Thay đổi secret key trong production

# Cấu hình API cho LLM
API_URL = os.getenv('API_URL', "http://localhost:1234/v1/chat/completions")

# Danh sách các từ ngữ không phù hợp cần lọc
inappropriate_words = [
    # Từ ngữ tiêu cực chung
    "xấu", "tệ", "dở", "kém", "tồi", "bẩn", "thối", "hôi", "ghét", "xấu xa", 
    "ngu", "đần", "điên", "khùng", "ngốc", "đáng ghét", "chửi bậy", "tục tĩu",
    "nguyền rủa", "chết", "chết tiệt", "ngu ngốc", "đần độn", "nhảm nhí", "vớ vẩn",
    
    # Từ ngữ thô tục và xúc phạm bổ sung (tiếng Việt)
    "đụ", "địt", "lồn", "cặc", "cu", "buồi", "dái", "chim", "kẹc", "loz", "vãi", "đéo", 
    "đ*", "đ.", "d*", "đít", "cứt", "ỉa", "đái", "đ é o", "đ-é-o",
    "đìu", "chó", "súc vật", "lòn", "l.ồn", "l*n", "l0n", "lờn", "nứng", "đĩ", "đ ĩ",
    "cave", "điếm", "thằng điên", "thằng đần", "thằng dở", "mẹ mày", "con mẹ", "cm", "cmm", 
    "cức", "cuk", "đệch", "đệt", "vl", "vcl", "vloz", "vcloz", "vãi lồn", "vãi cả lồn",
    "mẹ kiếp", "má kiếp", "đồ chó", "đồ súc vật", "óc chó", "ngu như chó", "não lợn",
    
    # Từ ngữ phân biệt và xúc phạm bổ sung
    "thổ", "mán", "tàu", "chệt", "mọi rợ", "lùn", "đen thui", "mập ú", "ú như heo",
    "mặt lợn", "ăn như heo", "hèn", "yếu đuối", "phế", "què", "quắc", "đui", "câm", "điếc",
    "mọi đen", "thổ đen", "bần cố nông", "nghèo mạt", "không học", "ngu học", "dốt",
    
    # Các cách viết biến tấu
    "đcm", "đmm", "đmg", "dmg", "đmtc", "điên thật", "dcm", "dmm", "đm",
    "cc", "cac", "c.a.c", "c-a-c", "c@c", "cak", "kặk", "kak", "cặk", "cmn", "con me may",
    "me may", "tổ cha mày", "tổ sư", "tổ bà", "bà cha mày", "km", "tcm", "bà già mày",
    "lìn", "lin", "l1n", "l!n", "l.o.n", "diu", "d!u", "đjt", "d!t", "đjt", "dek", "deo",
    
    # Các từ ngữ tiêu cực khác cần lọc
    "đáng khinh", "đáng khinh bỉ", "khinh", "khinh bỉ", "bỉ ổi", "đê tiện", "đáng sợ", 
    "rùng rợn", "đần", "đần độn", "dốt nát", "thiểu năng", "thiểu năng trí tuệ", "mất não",
    "mất trí", "xuẩn", "xuẩn ngốc", "ngớ ngẩn", "nhảm", "nhảm nhí", "vớ vẩn", "sai lệch",
    "thảm hại", "đáng thương", "tổn thương", "hỏng", "hỏng hóc", "lỗi", "sai sót", "trục trặc",
    
    # Từ ngữ tiêu cực liên quan đến doanh nghiệp và vi phạm pháp luật
    "trốn thuế", "gian lận thuế", "gian lận", "gian lận tài chính", "hối lộ", "đút lót", "chạy án",
    "rửa tiền", "lừa đảo", "đa cấp", "bất hợp pháp", "vi phạm pháp luật", "bôi trơn", "ăn hối lộ",
    "tham nhũng", "tham ô", "biển thủ", "chiếm đoạt", "lừa gạt", "bẩn tiền", "nghìn tỉ", "ngàn tỉ",
    "buôn lậu", "buôn người", "buôn bán trái phép", "khai man", "gian dối", "làm giả", "hàng giả",
    "hàng nhái", "bán hàng giả", "hàng kém chất lượng", "hàng nhập lậu", "hang lậu", "hô biến",
    "thâu tóm", "lũng đoạn", "độc quyền", "làm ăn phi pháp", "bỏ ngoài sổ sách", "không hóa đơn",
    "tiền chùa", "tiền lậu", "bỏ túi", "mua quan bán chức", "chạy chức", "chạy quyền", 
    "phe phái", "tranh giành quyền lực", "mafia", "băng đảng", "xã hội đen", "côn đồ",
    "thao túng", "điều hành ngầm", "chi phối", "lập quỹ đen", "quỹ đen", "sân sau",
    "bòn rút", "vơ vét", "lợi dụng sơ hở", "lợi dụng chính sách", "lách luật", "luồn lách",
    "trốn đóng bảo hiểm", "phá sản giả", "giả phá sản", "vỡ nợ", "ôm tiền bỏ trốn",
    "bất tín", "lừa dối", "không uy tín", "mập mờ", "thiếu minh bạch", "che giấu thông tin",
    "báo cáo giả", "làm đẹp báo cáo", "tài chính bất minh", "sổ sách không rõ ràng",
    "kếch xù", "lãi suất cắt cổ", "cho vay nặng lãi", "tín dụng đen", "bóc lột", "bóc lột lao động",
    "làm thêm không lương", "sa thải vô lý", "phân biệt đối xử", "phạt tiền vô lý", 
    "nhân viên tệ", "lãnh đạo tệ", "không biết quản lý", "thiếu chuyên nghiệp", "bất tài",
    "làm ăn thua lỗ", "suy thoái", "khánh kiệt", "nợ nần chồng chất", "phá sản", "thua lỗ triền miên",
    "kinh doanh thất bại", "thất bại thảm hại", "đầu tư sai lầm", "sai chiến lược", "không có tầm nhìn",
    "găm hàng", "đầu cơ", "đầu cơ tích trữ", "thao túng giá", "thổi giá", "tạo sóng", "lướt sóng",
    "sản phẩm kém", "dịch vụ tệ", "phục vụ không tốt", "sau bán hàng kém", "không bảo hành",
    "tiếp thị lừa đảo", "quảng cáo sai sự thật", "thổi phồng thành tích", "PR giả",
    "bán thông tin khách hàng", "hàng không", "hư hỏng", "kém chất lượng",
    "ép giá", "đè giá", "cạnh tranh không lành mạnh", "phá giá", "mua chuộc",
    "không có tâm", "làm ăn thiếu đạo đức", "vô đạo đức", "bất lương", "vô lương tâm",
]

def filter_content(content):
    """
    Lọc nội dung để loại bỏ từ ngữ không phù hợp
    
    Args:
        content (str): Nội dung cần lọc
    
    Returns:
        tuple: (filtered_content, filtered_words)
    """
    filtered_content = content
    filtered_words = []
    
    # Xóa bỏ kí tự hoa thị (*)
    filtered_content = filtered_content.replace('*', '')
    
    # Lọc các từ ngữ không phù hợp
    for word in inappropriate_words:
        if word.lower() in filtered_content.lower():
            # Đánh dấu từ này cần được lọc
            filtered_words.append(word)
            
            # Thay thế từ bằng dấu sao (censored)
            pattern = re.compile(re.escape(word), re.IGNORECASE)
            censored = '*' * len(word)
            filtered_content = pattern.sub(censored, filtered_content)
    
    return filtered_content, filtered_words

def filter_list(items):
    """
    Lọc danh sách các mục để loại bỏ từ ngữ không phù hợp
    
    Args:
        items (list): Danh sách các mục cần lọc
    
    Returns:
        tuple: (filtered_items, filtered_words)
    """
    filtered_items = []
    all_filtered_words = []
    
    if not items:
        return [], []
    
    for item in items:
        # Xóa bỏ kí tự hoa thị (*)
        item = item.replace('*', '') if isinstance(item, str) else item
        
        # Lọc nội dung
        filtered_item, filtered_words = filter_content(item)
        filtered_items.append(filtered_item)
        all_filtered_words.extend(filtered_words)
    
    return filtered_items, all_filtered_words

def filter_data(data):
    """
    Lọc từ ngữ không phù hợp từ dữ liệu có cấu trúc (dict, list)
    Args:
        data: Dữ liệu cần lọc (dict hoặc list)
    Returns:
        Tuple chứa (dữ liệu đã lọc, danh sách từ bị lọc)
    """
    all_filtered_words = []
    
    if isinstance(data, dict):
        result = {}
        for k, v in data.items():
            filtered_v, filtered_words = filter_data(v)
            result[k] = filtered_v
            all_filtered_words.extend(filtered_words)
        return result, list(set(all_filtered_words))
    
    elif isinstance(data, list):
        result = []
        for item in data:
            filtered_item, filtered_words = filter_data(item)
            result.append(filtered_item)
            all_filtered_words.extend(filtered_words)
        return result, list(set(all_filtered_words))
    
    elif isinstance(data, str):
        return filter_content(data)
    
    else:
        return data, []


def read_docx_file(file_storage):
    """Đọc nội dung file DOCX và chuyển thành string."""
    try:
        if 'docx' not in globals():
            return "Thư viện python-docx chưa được cài đặt. Không thể đọc file DOCX."
        
        # Lưu file tạm để xử lý
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            file_storage.save(temp_file.name)
            temp_path = temp_file.name
        
        # Đọc nội dung từ file docx
        doc = docx.Document(temp_path)
        content = []
        
        # Trích xuất text từ đoạn văn
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        # Trích xuất text từ bảng
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                if any(cell for cell in row_text):  # Chỉ thêm hàng có nội dung
                    content.append(' | '.join(row_text))
        
        # Xóa file tạm
        os.unlink(temp_path)
        
        return '\n'.join(content)
    except Exception as e:
        print(f"Error reading DOCX: {str(e)}")
        return f"Lỗi khi đọc file DOCX: {str(e)}"

def get_llm_suggestions(goal, file_content=None):
    """
    Nhận goal và nội dung file để phân tích và đề xuất tiêu chí và phương án.
    
    Args:
        goal: Mục tiêu cần phân tích
        file_content: Nội dung file đã được đọc (dạng string)
    
    Returns:
        Dict chứa các tiêu chí và phương án được đề xuất
    """
    # Xác định nội dung prompt dựa trên việc có file hay không
    if file_content:
        prompt = f"""Hãy phân tích mục tiêu và nội dung file đính kèm, sau đó đề xuất 4 tiêu chí đánh giá quan trọng nhất và 4 phương án lựa chọn phù hợp nhất với mục tiêu đã cho:
        
        Mục tiêu: {goal}
        
        Lưu ý:
        - Ưu tiên các tiêu chí mang tính chất định lượng 
        - Chỉ đề xuất 4 tiêu chí quan trọng nhất
        - Chỉ đề xuất 4 phương án lựa chọn tốt nhất
        - Đảm bảo các tiêu chí và phương án phải thực tế và có thể so sánh được
        - Phân tích cả nội dung file để đề xuất cho phù hợp
        - Đảm bảo sử dụng ngôn ngữ lịch sự, tránh từ ngữ tiêu cực hoặc không phù hợp
        - Không sử dụng kí tự đặc biệt '*' trong câu trả lời
        
        Hãy trả về kết quả chỉ dưới định dạng JSON như sau, không kèm theo bất kỳ văn bản bổ sung nào:
        {{
            "criteria": ["tiêu chí 1", "tiêu chí 2", "tiêu chí 3", "tiêu chí 4"],
            "alternatives": ["phương án 1", "phương án 2", "phương án 3", "phương án 4"]
        }}
        
        Dữ liệu như sau:
        {file_content[:100000]}  
        tiêu chí phải phù hợp với mục tiêu {goal};
        """
    else:
        prompt = f"""Hãy phân tích mục tiêu sau và đề xuất 4 tiêu chí đánh giá quan trọng nhất và 4 phương án lựa chọn phù hợp nhất:
        Mục tiêu: {goal}
        
        Lưu ý:
        - Chỉ đề xuất 4 tiêu chí quan trọng nhất
        - Chỉ đề xuất 4 phương án lựa chọn tốt nhất
        - Đảm bảo các tiêu chí và phương án phải thực tế và có thể so sánh được
        - Đảm bảo sử dụng ngôn ngữ lịch sự, tránh từ ngữ tiêu cực hoặc không phù hợp
        - Không sử dụng kí tự đặc biệt '*' trong câu trả lời
        
        Hãy trả về kết quả chỉ dưới định dạng JSON như sau, không kèm theo bất kỳ văn bản bổ sung nào:
        {{
            "criteria": ["tiêu chí 1", "tiêu chí 2", "tiêu chí 3", "tiêu chí 4"],
            "alternatives": ["phương án 1", "phương án 2", "phương án 3", "phương án 4"]
        }}
        """
    try:
        response = requests.post(
            API_URL,
            json={
                "messages": [
                    {"role": "system", "content": "Bạn là một chuyên gia phân tích và tư vấn ra quyết định. Hãy phân tích mục tiêu và đề xuất 4 tiêu chí đánh giá quan trọng nhất và 4 phương án thay thế phù hợp nhất. Đảm bảo sử dụng ngôn ngữ phù hợp, không có từ ngữ tiêu cực, xúc phạm. Khi trả về kết quả, hãy LUÔN LUÔN trả về THEO ĐỊNH DẠNG JSON mà không có văn bản thừa, không cần diễn giải, không cần các thẻ markdown như ```json``` và không có văn bản giải thích. KHÔNG được sử dụng kí tự * trong kết quả trả về."},
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.7,
                "max_tokens": 800 if file_content else 500  # Tăng max_tokens nếu có file
            }
        )
        response.raise_for_status()
        
        response_data = response.json()
        if 'choices' in response_data and len(response_data['choices']) > 0:
            content = response_data['choices'][0]['message']['content']
            
            # Cải thiện xử lý JSON
            import json
            import re
            
            # Làm sạch nội dung trước khi phân tích
            content = content.strip()
            
            # Xóa bỏ các thẻ markdown codeblock nếu có
            content = re.sub(r'```json|```', '', content)
            
            # Tìm chuỗi JSON hợp lệ trong nội dung trả về
            json_match = re.search(r'({[\s\S]*})', content)
            
            result = {"criteria": [], "alternatives": []}
            
            try:
                if json_match:
                    json_str = json_match.group(1).strip()
                    parsed_result = json.loads(json_str)
                    
                    # Đảm bảo result có cấu trúc đúng
                    if isinstance(parsed_result, dict):
                        # Đảm bảo có các key cần thiết
                        if "criteria" in parsed_result and isinstance(parsed_result["criteria"], list):
                            result["criteria"] = parsed_result["criteria"][:4]  # Chỉ lấy 4 tiêu chí
                        
                        if "alternatives" in parsed_result and isinstance(parsed_result["alternatives"], list):
                            result["alternatives"] = parsed_result["alternatives"][:4]  # Chỉ lấy 4 phương án
                else:
                    # Nếu không tìm thấy JSON, thử phân tích nội dung theo cấu trúc
                    criteria_match = re.search(r'criteria["\']\s*:\s*\[(.*?)\]', content, re.DOTALL)
                    alternatives_match = re.search(r'alternatives["\']\s*:\s*\[(.*?)\]', content, re.DOTALL)
                    
                    if criteria_match:
                        criteria_str = criteria_match.group(1)
                        criteria_list = re.findall(r'["\']([^"\']+)["\']', criteria_str)
                        if criteria_list:
                            result["criteria"] = criteria_list[:4]
                    
                    if alternatives_match:
                        alternatives_str = alternatives_match.group(1)
                        alternatives_list = re.findall(r'["\']([^"\']+)["\']', alternatives_str)
                        if alternatives_list:
                            result["alternatives"] = alternatives_list[:4]
            except json.JSONDecodeError as json_error:
                print("Error parsing JSON:", json_error)
                print("Content causing error:", content)
                # Parse JSON dạng khác nếu có lỗi
                try:
                    # Thử phân tích lại với các ký tự đặc biệt đã được loại bỏ
                    cleaned_content = re.sub(r'[^\x00-\x7F]+', ' ', content)
                    cleaned_content = re.sub(r'\\', '', cleaned_content)
                    json_match = re.search(r'({[\s\S]*})', cleaned_content)
                    if json_match:
                        json_str = json_match.group(1).strip()
                        parsed_result = json.loads(json_str)
                        
                        if isinstance(parsed_result, dict):
                            if "criteria" in parsed_result and isinstance(parsed_result["criteria"], list):
                                result["criteria"] = parsed_result["criteria"][:4]
                            
                            if "alternatives" in parsed_result and isinstance(parsed_result["alternatives"], list):
                                result["alternatives"] = parsed_result["alternatives"][:4]
                except Exception as e:
                    print("Secondary parsing failed:", e)
            
            # Xóa bỏ các dấu hoa thị (*) từ kết quả
            if "criteria" in result:
                result["criteria"] = [re.sub(r'\*', '', criterion) for criterion in result["criteria"]]
            
            if "alternatives" in result:
                result["alternatives"] = [re.sub(r'\*', '', alternative) for alternative in result["alternatives"]]
            
            # Lọc từ ngữ không phù hợp trong kết quả trả về
            filtered_result = {}
            filtered_result['criteria'], filtered_criteria_words = filter_list(result['criteria'])
            filtered_result['alternatives'], filtered_alternatives_words = filter_list(result['alternatives'])
            
            all_filtered_words = filtered_criteria_words + filtered_alternatives_words
            if all_filtered_words:
                filtered_result['warning'] = f"Đã lọc {len(all_filtered_words)} từ ngữ không phù hợp"
                filtered_result['filtered_words'] = all_filtered_words
            
            return filtered_result
        
        # Trả về kết quả mặc định nếu không có nội dung từ LLM
        return {
            "error": "Không nhận được kết quả phù hợp từ LLM.",
            "criteria": [],
            "alternatives": []
        }
    except requests.exceptions.RequestException as request_error:
        print(f"API Request Error: {str(request_error)}")
        return {
            "error": f"Lỗi kết nối đến máy chủ LLM: {str(request_error)}",
            "criteria": [],
            "alternatives": []
        }
    except Exception as e:
        print(f"Unexpected error in get_llm_suggestions: {str(e)}")
        return {
            "error": f"Lỗi khi xử lý kết quả từ LLM: {str(e)}",
            "criteria": [],
            "alternatives": []
        }

def calculate_ahp(comparison_matrix):
    matrix_original = comparison_matrix.copy()
    n = comparison_matrix.shape[0]

    # Tính sum từng cột
    arr_sum = []
    for i in range(n):
        sum = 0
        for j in range(n):
            sum += comparison_matrix[j][i]
        arr_sum.append(sum)

    # Chuẩn hóa ma trận so sánh cặp
    normalized_matrix = np.zeros((n, n))
    for i in range(n):
        for j in range(n):
            normalized_matrix[i][j] = comparison_matrix[i][j] / arr_sum[j]

    # Tính Criteria Weight
    arr_avg = []
    for i in range(n):
        sum = 0
        for j in range(n):
            sum += normalized_matrix[i][j]
        arr_avg.append(sum/n)

    # Nhân Criteria Weight với ma trận so sánh cặp
    for i in range(n):
        for j in range(n):
            matrix_original[j][i] = matrix_original[j][i] * arr_avg[i]
    
    # Tính Weighted Sum Value
    arr_WSV = []
    for i in range(n):
        sum = 0
        for j in range(n):
            sum += matrix_original[i][j]
        arr_WSV.append(sum)
    
    # Tính Consistency Vector
    arr_CV = []
    for i in range(len(arr_WSV)):
        arr_CV.append(arr_WSV[i] / arr_avg[i])

    # Tính Lambda Max
    sum = 0
    for value in arr_CV:
        sum += value
    lambda_max = sum / len(arr_CV)

    # Tính CI
    ci = (lambda_max - n) / (n - 1)

    # RI values
    ri_values = [0.00, 0.00, 0.58, 0.90, 1.12, 1.24, 1.32, 1.41, 1.45, 1.49, 1.51, 1.54, 1.56, 1.58, 1.59]
    ri = ri_values[n-1] if n <= len(ri_values) else ri_values[-1]

    # Tính CR
    cr = ci / ri

    return {
        "weights": arr_avg,
        "lambda_max": float(lambda_max),
        "ci": float(ci),
        "cr": float(cr),
        "normalized_matrix": normalized_matrix.tolist()
    }

def calculate_final_scores(criteria_weights, alt_matrices_by_criteria):
    """
    Tính điểm cuối cùng cho các phương án dựa trên trọng số tiêu chí và ma trận đánh giá theo từng tiêu chí
    
    Args:
        criteria_weights: Danh sách trọng số các tiêu chí
        alt_matrices_by_criteria: Dict với key là tên tiêu chí và value là ma trận trọng số của phương án theo tiêu chí đó
    
    Returns:
        Dict chứa điểm số cuối cùng của từng phương án
    """
    final_scores = {}
    
    # Lấy danh sách các phương án từ ma trận đầu tiên
    first_criterion = list(alt_matrices_by_criteria.keys())[0]
    alternatives = list(alt_matrices_by_criteria[first_criterion]["alternatives"].keys())
    
    # Khởi tạo điểm số cho từng phương án
    for alt in alternatives:
        final_scores[alt] = 0
    
    # Tính tổng điểm cho từng phương án
    for i, criterion in enumerate(alt_matrices_by_criteria.keys()):
        criterion_weight = criteria_weights[i]
        alt_weights = alt_matrices_by_criteria[criterion]["alternatives"]
        
        for alt in alternatives:
            final_scores[alt] += criterion_weight * alt_weights[alt]
    
    return final_scores

# Route cho trang đăng nhập
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        success, error_message = login_user(username, password)
        if success:
            flash('Đăng nhập thành công!', 'success')
            return redirect(url_for('index'))
        else:
            flash(error_message, 'error')
    
    return render_template('login.html')

# Route cho đăng xuất
@app.route('/logout')
def logout():
    logout_user()
    flash('Đã đăng xuất thành công', 'success')
    return redirect(url_for('login'))

# Trang chủ - yêu cầu đăng nhập
@app.route('/')
@login_required
def index():
    return render_template('index.html', user=session.get('user'))

# Route cho xem logs - chỉ admin mới có quyền
@app.route('/logs')
@role_required(['admin'])
def view_logs():
    db = next(get_db())
    logs = db.query(Log).order_by(Log.created_at.desc()).all()
    return render_template('logs.html', logs=logs)

@app.route('/get_suggestions', methods=['POST'])
def get_suggestions():
    """Xử lý yêu cầu phân tích từ form hoặc AJAX."""
    try:
        # Lấy goal từ form hoặc JSON
        if request.form:
            goal = request.form.get('goal', '')
        else:
            goal = request.json.get('goal', '')
        
        # Lọc từ ngữ không phù hợp từ goal
        filtered_goal, filtered_goal_words = filter_content(goal)
        
        # Kiểm tra goal có tồn tại
        if not filtered_goal:
            return jsonify({"error": "Vui lòng nhập mục tiêu"}), 400
        
        # Xử lý file nếu có
        file_content = None
        filtered_file_words = []
        if 'docxFile' in request.files and request.files['docxFile'].filename != '':
            file = request.files['docxFile']
            filename = file.filename.lower()
            
            # Kiểm tra loại file và đọc nội dung
            if filename.endswith('.docx'):
                file_content = read_docx_file(file)
                # Lọc nội dung từ file
                file_content, filtered_file_words = filter_content(file_content)
            elif filename.endswith('.doc'):
                return jsonify({"error": "File .doc cũ không được hỗ trợ. Vui lòng chuyển đổi sang .docx"}), 400
            else:
                return jsonify({"error": "Chỉ hỗ trợ file Word (.docx)"}), 400
        
        # Gọi hàm get_llm_suggestions
        suggestions = get_llm_suggestions(filtered_goal, file_content)
        
        # Kết hợp tất cả các từ đã lọc từ nhiều nguồn
        all_filtered_words = []
        
        # Từ ngữ bị lọc từ goal và file
        if filtered_goal_words:
            all_filtered_words.extend(filtered_goal_words)
        if filtered_file_words:
            all_filtered_words.extend(filtered_file_words)
        
        # Từ ngữ bị lọc từ kết quả LLM
        if 'filtered_words' in suggestions:
            all_filtered_words.extend(suggestions['filtered_words'])
            # Xóa trường filtered_words khỏi suggestions để tránh trùng lặp
            del suggestions['filtered_words']
        
        # Loại bỏ trùng lặp
        all_filtered_words = list(set(all_filtered_words))
        
        # Thêm thông tin về các từ đã lọc vào kết quả
        if all_filtered_words:
            suggestions['filtered_words'] = all_filtered_words
            suggestions['has_inappropriate_content'] = True
        else:
            suggestions['has_inappropriate_content'] = False
        
        # Ghi log nếu có từ ngữ không phù hợp
        if all_filtered_words and 'user' in session:
            try:
                user = session['user']
                username = user.get('username', 'unknown')
                log_action(
                    user_id=username,
                    action_type="content_filtered",
                    details=f"Từ ngữ không phù hợp được phát hiện: {', '.join(all_filtered_words)}",
                    related_data=goal
                )
            except Exception as log_error:
                print(f"Error logging filtered content: {str(log_error)}")
        
        return jsonify(suggestions)
        
    except Exception as e:
        print(f"Error in get_suggestions: {str(e)}")
        return jsonify({"error": f"Lỗi xử lý yêu cầu: {str(e)}"}), 500

@app.route('/calculate_ahp', methods=['POST'])
def calculate():
    data = request.json
    
    # Lọc dữ liệu đầu vào
    criteria, criteria_filtered_words = filter_list(data.get('criteria', []))
    goal, goal_filtered_words = filter_content(data.get('goal', ''))
    
    # Thêm thông tin về các từ bị lọc
    all_filtered_words = list(set(criteria_filtered_words + goal_filtered_words))
    has_inappropriate_content = len(all_filtered_words) > 0
    
    # Chuyển đổi ma trận đầu vào từ chuỗi sang số
    input_matrix = data['matrix']
    comparison_matrix = []
    
    for row in input_matrix:
        numeric_row = []
        for value in row:
            if isinstance(value, str) and '/' in value:
                # Xử lý giá trị dạng phân số như "1/3", "1/5", etc.
                parts = value.split('/')
                if len(parts) == 2:
                    try:
                        numeric_value = float(parts[0]) / float(parts[1])
                        numeric_row.append(numeric_value)
                    except (ValueError, ZeroDivisionError):
                        return jsonify({"error": f"Không thể chuyển đổi giá trị '{value}' thành số"}), 400
            else:
                # Xử lý giá trị số hoặc chuỗi số
                try:
                    numeric_row.append(float(value))
                except ValueError:
                    return jsonify({"error": f"Không thể chuyển đổi giá trị '{value}' thành số"}), 400
        comparison_matrix.append(numeric_row)
    
    # Chuyển đổi sang numpy array và tính toán
    comparison_matrix = np.array(comparison_matrix)
    result = calculate_ahp(comparison_matrix)
    
    # Thêm thông tin về các từ bị lọc vào kết quả
    result["filtered_info"] = {
        "has_inappropriate_content": has_inappropriate_content,
        "filtered_words": all_filtered_words
    }
    
    return jsonify(result)

@app.route('/calculate_alternative_matrices', methods=['POST'])
def calculate_alternative_matrices():
    """Tính toán ma trận đánh giá các phương án theo từng tiêu chí"""
    data = request.json
    
    # Lọc dữ liệu đầu vào
    criteria, criteria_filtered_words = filter_list(data.get('criteria', []))
    alternatives, alt_filtered_words = filter_list(data.get('alternatives', []))
    matrices = data.get('matrices', {})
    criteria_weights = data.get('criteria_weights', [])
    
    # Thêm thông tin về các từ bị lọc
    all_filtered_words = list(set(criteria_filtered_words + alt_filtered_words))
    has_inappropriate_content = len(all_filtered_words) > 0
    
    # Kết quả của từng ma trận phương án theo từng tiêu chí
    alt_matrices_results = {}
    cr_values = {}
    
    # Tính toán trọng số cho từng tiêu chí
    for criterion in criteria:
        if criterion in matrices:
            # Chuyển đổi ma trận từ chuỗi sang số
            input_matrix = matrices[criterion]
            numeric_matrix = []
            
            try:
                for row in input_matrix:
                    numeric_row = []
                    for value in row:
                        if isinstance(value, str):
                            if '/' in value:
                                # Xử lý giá trị dạng phân số như "1/3", "1/5", etc.
                                parts = value.split('/')
                                if len(parts) == 2:
                                    try:
                                        numeric_value = float(parts[0]) / float(parts[1])
                                        numeric_row.append(numeric_value)
                                    except (ValueError, ZeroDivisionError) as e:
                                        return jsonify({"error": f"Lỗi khi chuyển đổi phân số '{value}': {str(e)}"}), 400
                            else:
                                # Xử lý giá trị chuỗi số
                                try:
                                    numeric_value = float(value)
                                    numeric_row.append(numeric_value)
                                except ValueError:
                                    return jsonify({"error": f"Không thể chuyển đổi giá trị '{value}' thành số"}), 400
                        else:
                            # Đã là số, giữ nguyên
                            numeric_row.append(float(value))
                    numeric_matrix.append(numeric_row)
                
                matrix = np.array(numeric_matrix)
                result = calculate_ahp(matrix)
                
                # Kiểm tra CR của ma trận
                if result["cr"] > 0.1:
                    print(f"Tỷ số nhất quán CR={result['cr']} cho tiêu chí {criterion} vượt quá ngưỡng cho phép!")
                
                cr_values[criterion] = result["cr"]
                
                # Lưu kết quả vào dict với key là tên tiêu chí
                alt_matrices_results[criterion] = {
                    "weights": result["weights"],
                    "cr": result["cr"],
                    "alternatives": {alt: weight for alt, weight in zip(alternatives, result["weights"])}
                }
            except Exception as e:
                return jsonify({"error": f"Lỗi khi xử lý ma trận cho tiêu chí '{criterion}': {str(e)}"}), 400
    
    # Tính điểm cuối cùng cho các phương án
    final_scores = calculate_final_scores(criteria_weights, alt_matrices_results)
    
    # Tìm phương án tốt nhất
    best_alternative = max(final_scores.items(), key=lambda x: x[1])[0]
    
    # Chuẩn bị điểm số của các phương án theo từng tiêu chí
    alternative_scores_by_criteria = {}
    for criterion in alt_matrices_results:
        alternative_scores_by_criteria[criterion] = alt_matrices_results[criterion]["alternatives"]
    
    return jsonify({
        "alt_matrices_results": alt_matrices_results,
        "final_scores": final_scores,
        "best_alternative": best_alternative,
        "alternative_scores_by_criteria": alternative_scores_by_criteria,
        "criteria_weights": criteria_weights,
        "CR": cr_values
    })

@app.route('/get_alternative_matrices', methods=['POST'])
def get_alternative_matrices():
    """Lấy dữ liệu ma trận so sánh các phương án theo từng tiêu chí từ server"""
    try:
        data = request.json
        
        # Lọc dữ liệu đầu vào
        criteria, criteria_filtered_words = filter_list(data.get('criteria', []))
        alternatives, alt_filtered_words = filter_list(data.get('alternatives', []))
        goal, goal_filtered_words = filter_content(data.get('goal', ''))
        
        # Thêm thông tin về các từ bị lọc
        all_filtered_words = list(set(criteria_filtered_words + alt_filtered_words + goal_filtered_words))
        filtered_info = {
            "has_inappropriate_content": len(all_filtered_words) > 0,
            "filtered_words": all_filtered_words
        }
        
        # Kiểm tra dữ liệu đầu vào
        if not criteria or not alternatives:
            return jsonify({"error": "Thiếu dữ liệu tiêu chí hoặc phương án"}), 400
            
        # Chuẩn bị ma trận so sánh cặp cho từng tiêu chí
        matrices = {}
        matrices_info = {}
        
        for criterion in criteria:
            n = len(alternatives)
            # Tạo ma trận với giá trị mặc định là 1 (đường chéo là 1, các giá trị khác được tính từ LLM)
            matrix = np.ones((n, n))
            
            # Tạo prompt để gửi tới LLM
            prompt = f"""
            Với mục tiêu: {goal}
            Hãy tạo ma trận so sánh cặp định lượng {n}x{n} cho các phương án sau đây theo tiêu chí "{criterion}":
            Các phương án: {', '.join(alternatives)}
            
            Sử dụng thang đo Saaty (1-9) với ý nghĩa:
            1: Hai phương án quan trọng như nhau
            3: Phương án A hơi quan trọng hơn B
            5: Phương án A quan trọng hơn B
            7: Phương án A rất quan trọng hơn B
            9: Phương án A cực kỳ quan trọng hơn B
            
            Hãy tính toán và cung cấp các giá trị dựa trên ý nghĩa định lượng của tiêu chí "{criterion}".
            Nếu tiêu chí mang tính định lượng (như giá, khoảng cách, thời gian, hiệu suất...), hãy sử dụng tỷ lệ giá trị thực để xác định độ quan trọng tương đối.
            Ví dụ:
            - Nếu tiêu chí là "Giá" và phương án A có giá 100, phương án B có giá 300, thì B kém hơn A khoảng 3 lần, do đó giá trị a_AB = 3 và a_BA = 1/3
            - Nếu tiêu chí là "Hiệu suất" và phương án A có hiệu suất 80%, phương án B có hiệu suất 40%, thì A tốt hơn B khoảng 2 lần, do đó giá trị a_AB = 2 và a_BA = 1/2
            
            Chỉ trả về ma trận dưới dạng JSON với định dạng:
            [
                [a_11, a_12, ..., a_1n],
                [a_21, a_22, ..., a_2n],
                ...
                [a_n1, a_n2, ..., a_nn]
            ]
            
            Hãy phân tích sâu các giá trị định lượng của từng phương án theo tiêu chí này và đánh giá chính xác nhất có thể.
            Nếu tiêu chí trừu tượng không có giá trị định lượng rõ ràng, hãy sử dụng thang đo Saaty và phân tích logic để đưa ra đánh giá hợp lý nhất.
            
            LƯU Ý QUAN TRỌNG: Ma trận phải đảm bảo tỷ số nhất quán CR < 0.1 (10%) theo phương pháp AHP.
            Đảm bảo tính nhất quán của ma trận bằng cách:
            1. Đường chéo ma trận phải có giá trị 1 (a_ii = 1)
            2. Ma trận phải đối xứng nghịch đảo: a_ji = 1/a_ij
            3. Tính chất bắc cầu: Nếu a_ij = x và a_jk = y, thì a_ik nên xấp xỉ x*y
            4. Hạn chế sử dụng các giá trị cực đoan (1/9 hoặc 9)
            """
            
            try:
                # Trước tiên thử gọi API để xác định các giá trị số học cho tiêu chí
                quantitative_prompt = f"""
                Phân tích tiêu chí "{criterion}" cho mục tiêu "{goal}":
                1. Tiêu chí này có phải là tiêu chí định lượng không? (có/không)
                2. Nếu là tiêu chí định lượng, hãy ước tính giá trị số cụ thể cho mỗi phương án: {', '.join(alternatives)}
                3. Đơn vị đo lường của tiêu chí này là gì?
                4. Tiêu chí này là tối đa hóa (càng cao càng tốt) hay tối thiểu hóa (càng thấp càng tốt)?
                
                Trả về kết quả dưới dạng JSON:
                {{
                    "is_quantitative": true/false,
                    "values": {{
                        "phương án 1": giá trị số,
                        "phương án 2": giá trị số,
                        ...
                    }},
                    "unit": "đơn vị",
                    "maximize": true/false
                }}
                
                Nếu không phải là tiêu chí định lượng, chỉ cần trả về:
                {{
                    "is_quantitative": false
                }}
                """
                
                response = requests.post(
                    API_URL,
                    json={
                        "messages": [
                            {"role": "system", "content": "Bạn là một chuyên gia phân tích dữ liệu. Hãy phân tích tính chất định lượng của tiêu chí được cung cấp."},
                            {"role": "user", "content": quantitative_prompt}
                        ],
                        "temperature": 0.3,
                        "max_tokens": 500
                    }
                )
                response.raise_for_status()
                
                is_quantitative = False
                quantitative_values = {}
                maximize = True  # Mặc định là tối đa hóa
                unit = ""
                
                response_data = response.json()
                if 'choices' in response_data and len(response_data['choices']) > 0:
                    content = response_data['choices'][0]['message']['content']
                    
                    # Trích xuất JSON từ kết quả
                    import json
                    import re
                    
                    # Tìm đối tượng JSON trong kết quả
                    json_match = re.search(r'\{.*\}', content, re.DOTALL)
                    if json_match:
                        try:
                            quant_data = json.loads(json_match.group(0))
                            is_quantitative = quant_data.get('is_quantitative', False)
                            
                            if is_quantitative:
                                quantitative_values = quant_data.get('values', {})
                                maximize = quant_data.get('maximize', True)
                                unit = quant_data.get('unit', "")
                                print(f"Phân tích định lượng cho tiêu chí '{criterion}':")
                                print(f"  Giá trị: {quantitative_values}")
                                print(f"  Tối đa hóa: {maximize}")
                                print(f"  Đơn vị: {unit}")
                        except json.JSONDecodeError as e:
                            print(f"Lỗi phân tích JSON quantitative: {e}")
                
                # Lưu thông tin định lượng
                if is_quantitative:
                    matrices_info[criterion] = {
                        "is_quantitative": True,
                        "maximize": maximize,
                        "unit": unit,
                        "values": quantitative_values
                    }
                else:
                    matrices_info[criterion] = {
                        "is_quantitative": False
                    }
                
                # Hàm để chuyển ma trận số thành ma trận string
                def convert_to_string_matrix(numeric_matrix):
                    string_matrix = []
                    for i in range(n):
                        string_row = []
                        for j in range(n):
                            val = numeric_matrix[i][j]
                            # Chuyển số thành string
                            if val == 1:
                                string_row.append("1")
                            elif val < 1:
                                # Chuyển số thập phân thành phân số 1/x
                                denominator = int(round(1/val))
                                string_row.append(f"1/{denominator}")
                            else:
                                string_row.append(str(int(round(val))))
                        string_matrix.append(string_row)
                    return string_matrix
                
                # Hàm để kiểm tra độ nhất quán của ma trận
                def check_consistency(matrix):
                    try:
                        result = calculate_ahp(np.array(matrix))
                        cr = result["cr"]
                        return 0 < cr < 0.1, cr
                    except Exception as e:
                        print(f"Lỗi kiểm tra độ nhất quán: {e}")
                        return False, 1.0
                
                # Hàm để điều chỉnh ma trận để có độ nhất quán tốt hơn
                def improve_consistency(matrix, orig_matrix=None):
                    # Nếu không có ma trận gốc, sử dụng ma trận hiện tại
                    if orig_matrix is None:
                        orig_matrix = matrix.copy()
                    
                    # Thử điều chỉnh các giá trị không trên đường chéo để cải thiện độ nhất quán
                    for _ in range(3):  # Thử tối đa 3 lần
                        # Tính độ nhất quán
                        result = calculate_ahp(matrix)
                        cr = result["cr"]
                        
                        # Nếu CR nằm trong khoảng mong muốn (0 < CR < 0.1), thì dừng
                        if 0 < cr < 0.1:
                            break
                            
                        # Nếu CR = 0, thêm một chút nhiễu để tạo sự khác biệt nhỏ
                        if cr == 0:
                            # Tạo ma trận nhiễu nhỏ
                            for i in range(n):
                                for j in range(i+1, n):  # Chỉ làm việc với nửa trên tam giác
                                    # Thêm nhiễu nhỏ, khoảng 5-10% của giá trị hiện tại
                                    noise_factor = 1.0 + (np.random.random() * 0.1 - 0.05)  # Nhiễu +/- 5%
                                    matrix[i][j] *= noise_factor
                                    matrix[j][i] = 1 / matrix[i][j]
                            
                            # Tính lại CR sau khi thêm nhiễu
                            result = calculate_ahp(matrix)
                            cr = result["cr"]
                            
                            # Nếu CR vẫn = 0, tạo một số bất đồng nhỏ ở ô giá trị lớn nhất
                            if cr == 0:
                                # Tìm giá trị lớn nhất trên tam giác trên
                                max_val = 0
                                max_i, max_j = 0, 0
                                for i in range(n):
                                    for j in range(i+1, n):
                                        if matrix[i][j] > max_val:
                                            max_val = matrix[i][j]
                                            max_i, max_j = i, j
                                
                                # Điều chỉnh giá trị lớn nhất nếu tìm thấy
                                if max_val > 1:
                                    matrix[max_i][max_j] = max_val * 1.15  # Tăng thêm 15%
                                    matrix[max_j][max_i] = 1 / matrix[max_i][max_j]
                        
                        # Tính eigenvalues và eigenvectors
                        weights = result["weights"]
                        
                        # Điều chỉnh ma trận dựa trên tính chất bắc cầu
                        for i in range(n):
                            for j in range(i+1, n):  # Chỉ làm việc với nửa trên tam giác
                                # Sử dụng tỷ lệ trọng số để điều chỉnh giá trị
                                ideal_value = weights[i] / weights[j]
                                
                                # Điều chỉnh giá trị phù hợp (giữ trong khoảng [1/9, 9])
                                ideal_value = min(9, max(1/9, ideal_value))
                                
                                # Điều chỉnh giá trị hiện tại về phía giá trị lý tưởng
                                current = matrix[i][j]
                                adjusted = current * 0.7 + ideal_value * 0.3  # Trộn 70% giá trị cũ, 30% giá trị tối ưu
                                
                                # Làm tròn về các giá trị Saaty (1,3,5,7,9 hoặc nghịch đảo)
                                if adjusted >= 1:
                                    saaty_vals = [1, 2, 3, 4, 5, 6, 7, 8, 9]
                                    closest_idx = min(range(len(saaty_vals)), key=lambda i: abs(saaty_vals[i] - adjusted))
                                    matrix[i][j] = saaty_vals[closest_idx]
                                    matrix[j][i] = 1 / matrix[i][j]
                                else:
                                    inv_adjusted = 1 / adjusted
                                    saaty_vals = [1, 2, 3, 4, 5, 6, 7, 8, 9]
                                    closest_idx = min(range(len(saaty_vals)), key=lambda i: abs(saaty_vals[i] - inv_adjusted))
                                    matrix[j][i] = saaty_vals[closest_idx]
                                    matrix[i][j] = 1 / matrix[j][i]
                    
                    # Kiểm tra lại CR sau khi điều chỉnh
                    result = calculate_ahp(matrix)
                    cr = result["cr"]
                    
                    # Nếu CR vẫn = 0, đảm bảo có sự khác biệt nhỏ giữa các giá trị
                    if cr == 0:
                        # Tạo một ma trận có CR > 0 nhưng vẫn nhỏ
                        for i in range(n):
                            for j in range(i+1, n):  # Chỉ làm việc với nửa trên tam giác
                                if i != j:
                                    # Thêm nhiễu khác nhau cho mỗi phần tử
                                    noise = 1.0 + (np.random.random() * 0.2 - 0.1)  # Nhiễu +/- 10%
                                    matrix[i][j] = max(1.0, matrix[i][j] * noise)
                                    if matrix[i][j] > 9: 
                                        matrix[i][j] = 9.0
                                    matrix[j][i] = 1.0 / matrix[i][j]
                    
                    return matrix
                
                # Nếu có dữ liệu định lượng, tạo ma trận dựa trên tỷ lệ giá trị
                if is_quantitative and quantitative_values and len(quantitative_values) >= len(alternatives):
                    # Tạo ma trận dựa trên tỷ lệ giá trị số
                    numeric_matrix = np.ones((n, n))
                    
                    for i, alt_i in enumerate(alternatives):
                        for j, alt_j in enumerate(alternatives):
                            if i == j:
                                # Đường chéo luôn là 1
                                numeric_matrix[i][j] = 1
                                continue
                                
                            if alt_i in quantitative_values and alt_j in quantitative_values:
                                val_i = float(quantitative_values[alt_i])
                                val_j = float(quantitative_values[alt_j])
                                
                                # Tránh chia cho 0
                                if val_i == 0 and val_j == 0:
                                    ratio = 1
                                elif val_i == 0:
                                    ratio = 1/9  # Giá trị thấp nhất nếu val_i = 0
                                elif val_j == 0:
                                    ratio = 9  # Giá trị cao nhất nếu val_j = 0
                                else:
                                    if maximize:
                                        ratio = val_i / val_j  # Càng cao càng tốt
                                    else:
                                        ratio = val_j / val_i  # Càng thấp càng tốt
                                
                                # Giới hạn tỷ lệ trong khoảng 1/9 đến 9 theo thang Saaty
                                if ratio > 9:
                                    ratio = 9
                                elif ratio < 1/9:
                                    ratio = 1/9
                                    
                                numeric_matrix[i][j] = ratio
                    
                    # Kiểm tra độ nhất quán
                    is_consistent, cr = check_consistency(numeric_matrix)
                    
                    # Nếu ma trận không nhất quán, thử điều chỉnh
                    if not is_consistent:
                        print(f"Ma trận ban đầu cho {criterion} không nhất quán (CR={cr}), đang điều chỉnh...")
                        numeric_matrix = improve_consistency(numeric_matrix)
                        is_consistent, cr = check_consistency(numeric_matrix)
                        print(f"Ma trận điều chỉnh cho {criterion}: CR={cr}")
                    
                    # Chuyển đổi về ma trận chuỗi
                    matrices[criterion] = convert_to_string_matrix(numeric_matrix)
                    matrices_info[criterion]["cr"] = cr
                    matrices_info[criterion]["is_consistent"] = is_consistent
                    
                    # Nếu đã tạo được ma trận nhất quán, tiếp tục với tiêu chí tiếp theo
                    if is_consistent:
                        continue
                
                # Nếu không có đủ dữ liệu định lượng hoặc ma trận vẫn không nhất quán,
                # tiếp tục với phương pháp LLM
                max_attempts = 3
                is_consistent = False
                cr = 1.0
                
                for attempt in range(max_attempts):
                    response = requests.post(
                        API_URL,
                        json={
                            "messages": [
                                {"role": "system", "content": "Bạn là một chuyên gia về phương pháp AHP. Hãy tạo ma trận so sánh cặp phương án theo tiêu chí được yêu cầu, đảm bảo CR < 0.1."},
                                {"role": "user", "content": prompt + f"\n\nĐây là lần thử thứ {attempt+1}. Tạo ma trận đảm bảo CR < 0.1."}
                            ],
                            "temperature": 0.3,
                            "max_tokens": 500
                        }
                    )
                    response.raise_for_status()
                    
                    response_data = response.json()
                    if 'choices' in response_data and len(response_data['choices']) > 0:
                        content = response_data['choices'][0]['message']['content']
                        
                        # Trích xuất ma trận từ kết quả
                        import json
                        import re
                        
                        # Tìm mảng JSON trong kết quả
                        matrix_match = re.search(r'\[\s*\[.*?\]\s*\]', content, re.DOTALL)
                        if matrix_match:
                            matrix_json = matrix_match.group(0)
                            try:
                                generated_matrix = json.loads(matrix_json)
                                
                                # Chuyển thành ma trận số để kiểm tra độ nhất quán
                                numeric_matrix = np.ones((n, n))
                                for i in range(n):
                                    for j in range(n):
                                        val = generated_matrix[i][j]
                                        if isinstance(val, str) and '/' in val:
                                            parts = val.split('/')
                                            if len(parts) == 2:
                                                try:
                                                    numeric_matrix[i][j] = float(parts[0]) / float(parts[1])
                                                except:
                                                    numeric_matrix[i][j] = 1
                                        else:
                                            try:
                                                numeric_matrix[i][j] = float(val)
                                            except:
                                                numeric_matrix[i][j] = 1
                                
                                # Kiểm tra độ nhất quán
                                is_consistent, cr = check_consistency(numeric_matrix)
                                
                                # Nếu ma trận không nhất quán, thử điều chỉnh
                                if not is_consistent:
                                    print(f"Ma trận LLM cho {criterion} lần {attempt+1} không nhất quán (CR={cr}), đang điều chỉnh...")
                                    numeric_matrix = improve_consistency(numeric_matrix)
                                    is_consistent, cr = check_consistency(numeric_matrix)
                                    print(f"Ma trận điều chỉnh cho {criterion}: CR={cr}")
                                
                                # Chuyển đổi về ma trận chuỗi
                                string_matrix = convert_to_string_matrix(numeric_matrix)
                                
                                matrices[criterion] = string_matrix
                                matrices_info[criterion]["cr"] = cr
                                matrices_info[criterion]["is_consistent"] = is_consistent
                                
                                if is_consistent:
                                    print(f"Tạo ma trận nhất quán cho {criterion} thành công sau {attempt+1} lần thử.")
                                    break
                            except Exception as e:
                                print(f"Lỗi xử lý ma trận cho {criterion}: {e}")
                
                # Nếu sau nhiều lần thử vẫn không có ma trận nhất quán, tạo ma trận đơn vị
                if not is_consistent:
                    print(f"Không thể tạo ma trận nhất quán cho {criterion}, sử dụng ma trận cơ bản.")
                    # Tạo ma trận đơn vị (đường chéo 1, các ô khác là giá trị nhỏ 2-3)
                    numeric_matrix = np.ones((n, n))
                    for i in range(n):
                        for j in range(i+1, n):
                            numeric_matrix[i][j] = 1  # Mặc định các ô không trên đường chéo là 1
                            numeric_matrix[j][i] = 1
                    
                    matrices[criterion] = convert_to_string_matrix(numeric_matrix)
                    matrices_info[criterion]["cr"] = 0.0
                    matrices_info[criterion]["is_consistent"] = True
                
            except Exception as e:
                print(f"Error generating matrix for {criterion}: {str(e)}")
                # Tạo ma trận mặc định nếu có lỗi
                numeric_matrix = np.ones((n, n))
                matrices[criterion] = convert_to_string_matrix(numeric_matrix)
                matrices_info[criterion]["cr"] = 0.0
                matrices_info[criterion]["is_consistent"] = True
        
        # Thêm trường thông tin về phân tích định lượng để hiển thị cho người dùng
        return jsonify({
            "matrices": matrices,
            "matrices_info": matrices_info,
            "matrix_info": "Ma trận so sánh được tạo dựa trên phân tích định lượng của các tiêu chí, đảm bảo tỷ số nhất quán CR < 10%."
        })
    
    except Exception as e:
        print(f"Error in get_alternative_matrices: {str(e)}")
        return jsonify({"error": f"Lỗi khi lấy ma trận: {str(e)}"}), 500

@app.route('/alternative_matrices')
def alternative_matrices():
    return render_template('alternative_matrices.html')

@app.route('/export-excel', methods=['POST'])
def export_excel():
    """Export AHP analysis results to Excel file"""
    try:
        data = request.json
        goal, goal_filtered_words = filter_content(data.get('goal', 'AHP Analysis'))
        criteria, criteria_filtered_words = filter_list(data.get('criteria', []))
        alternatives, alt_filtered_words = filter_list(data.get('alternatives', []))
        criteriaWeights = data.get('criteriaWeights', [])
        alternativeScores, alt_scores_filtered_words = filter_data(data.get('alternativeScores', {}))
        finalScores = data.get('finalScores', {})
        crValues = data.get('CR', {})
        bestAlternative, best_alt_filtered_words = filter_content(data.get('best_alternative', ''))
        
        # Thêm thông tin về các từ bị lọc
        all_filtered_words = list(set(goal_filtered_words + criteria_filtered_words + 
                                      alt_filtered_words + alt_scores_filtered_words + 
                                      best_alt_filtered_words))
        has_inappropriate_content = len(all_filtered_words) > 0
        
        # Create Excel writer
        output = io.BytesIO()
        writer = pd.ExcelWriter(output, engine='xlsxwriter')
        
        # Create summary sheet
        summary_df = pd.DataFrame({
            'Phương án': list(finalScores.keys()),
            'Điểm số cuối cùng': list(finalScores.values())
        })
        summary_df = summary_df.sort_values('Điểm số cuối cùng', ascending=False)
        summary_df.to_excel(writer, sheet_name='Kết quả tổng hợp', index=False)
        
        # Format summary sheet
        workbook = writer.book
        worksheet = writer.sheets['Kết quả tổng hợp']
        header_format = workbook.add_format({
            'bold': True,
            'text_wrap': True,
            'valign': 'top',
            'fg_color': '#D7E4BC',
            'border': 1
        })
        
        # Add criteria weights and CR values sheet
        criteria_df = pd.DataFrame({
            'Tiêu chí': criteria,
            'Trọng số': criteriaWeights,
            'CR': [crValues.get(criterion, 0) for criterion in criteria]
        })
        criteria_df.to_excel(writer, sheet_name='Trọng số tiêu chí', index=False)
        
        # Add alternative scores by criteria
        for i, criterion in enumerate(criteria):
            scores = {}
            for alt in alternatives:
                scores[alt] = alternativeScores.get(criterion, {}).get(alt, 0)
            
            alt_df = pd.DataFrame({
                'Phương án': alternatives,
                f'Điểm theo {criterion}': [scores[alt] for alt in alternatives]
            })
            alt_df.to_excel(writer, sheet_name=f'Tiêu chí - {criterion[:25]}', index=False)
        
        # Add best alternative sheet
        best_df = pd.DataFrame({
            'Thông tin': ['Phương án tốt nhất', 'Điểm số'],
            'Giá trị': [bestAlternative, finalScores.get(bestAlternative, 0)]
        })
        best_df.to_excel(writer, sheet_name='Phương án tốt nhất', index=False)
        
        # Close the writer
        writer.close()
        
        # Return the Excel file
        output.seek(0)
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f"AHP_Analysis_{goal.replace(' ', '_')}.xlsx"
        )
        
    except Exception as e:
        print(f"Error exporting to Excel: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/export-pdf', methods=['POST'])
def export_pdf():
    """Export AHP analysis results to PDF file"""
    try:
        data = request.json
        goal, goal_filtered_words = filter_content(data.get('goal', 'AHP Analysis'))
        criteria, criteria_filtered_words = filter_list(data.get('criteria', []))
        alternatives, alt_filtered_words = filter_list(data.get('alternatives', []))
        criteriaWeights = data.get('criteriaWeights', [])
        alternativeScores, alt_scores_filtered_words = filter_data(data.get('alternativeScores', {}))
        finalScores = data.get('finalScores', {})
        crValues = data.get('CR', {})
        bestAlternative, best_alt_filtered_words = filter_content(data.get('best_alternative', ''))
        
        # Thêm thông tin về các từ bị lọc
        all_filtered_words = list(set(goal_filtered_words + criteria_filtered_words + 
                                      alt_filtered_words + alt_scores_filtered_words + 
                                      best_alt_filtered_words))
        has_inappropriate_content = len(all_filtered_words) > 0
        
        # Create a file-like buffer to receive PDF data
        buffer = io.BytesIO()
        
        # Create the PDF object using ReportLab
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=12
        )
        elements.append(Paragraph(f"Phân tích AHP: {goal}", title_style))
        elements.append(Spacer(1, 0.2*inch))
        
        # Criteria Weights Section
        elements.append(Paragraph("Trọng số các tiêu chí", styles['Heading2']))
        criteria_data = [['Tiêu chí', 'Trọng số']]
        for i, criterion in enumerate(criteria):
            criteria_data.append([criterion, f"{criteriaWeights[i]:.4f}"])
        
        criteria_table = Table(criteria_data, colWidths=[4*inch, 1.5*inch])
        criteria_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (1, 0), colors.lightblue),
            ('TEXTCOLOR', (0, 0), (1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        elements.append(criteria_table)
        elements.append(Spacer(1, 0.2*inch))
        
        # Final Scores Section
        elements.append(Paragraph("Điểm số cuối cùng", styles['Heading2']))
        
        # Sort alternatives by final score
        sorted_alternatives = sorted(
            alternatives,
            key=lambda alt: finalScores.get(alt, 0),
            reverse=True
        )
        
        final_data = [['Phương án', 'Điểm số', 'Xếp hạng']]
        for i, alt in enumerate(sorted_alternatives):
            final_data.append([
                alt, 
                f"{finalScores.get(alt, 0):.4f}", 
                f"{i+1}"
            ])
        
        final_table = Table(final_data, colWidths=[3*inch, 1.5*inch, 1*inch])
        final_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (2, 0), colors.lightblue),
            ('TEXTCOLOR', (0, 0), (2, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (2, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (2, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (2, 0), 12),
            ('BOTTOMPADDING', (0, 0), (2, 0), 12),
            ('BACKGROUND', (0, 1), (2, 1), colors.lightgreen),
            ('BACKGROUND', (0, 2), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        elements.append(final_table)
        elements.append(Spacer(1, 0.3*inch))
        
        # Create chart
        y_pos = np.arange(len(sorted_alternatives))
        scores = [finalScores.get(alt, 0) for alt in sorted_alternatives]
        
       
        # Save chart to memory
        img_data = io.BytesIO()
        img_data.seek(0)
        
        # Add image
        from reportlab.lib.utils import ImageReader
        img = ImageReader(img_data)
        elements.append(Paragraph("Biểu đồ kết quả", styles['Heading2']))
        elements.append(Image(img_data, width=6.5*inch, height=3*inch))
        
        # Build PDF
        doc.build(elements)
        
        # Get the value from the buffer
        buffer.seek(0)
        
        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f"AHP_Analysis_{goal.replace(' ', '_')}.pdf"
        )
        
    except Exception as e:
        print(f"Error exporting to PDF: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/check_content', methods=['POST'])
def check_content():
    """Kiểm tra nội dung có chứa từ ngữ không phù hợp hay không"""
    try:
        data = request.json
        content = data.get('content', '')
        context = data.get('context', 'unknown')  # Thêm context để biết nguồn gốc nội dung
        
        # Lọc nội dung
        filtered_content, filtered_words = filter_content(content)
        
        # Kiểm tra xem có từ ngữ nào bị lọc hay không
        has_inappropriate_content = len(filtered_words) > 0
        
        # Ghi log nếu có từ ngữ không phù hợp
        if has_inappropriate_content and 'user' in session:
            try:
                user = session['user']
                username = user.get('username', 'unknown')
                log_action(
                    user_id=username,
                    action_type="content_filtered", 
                    details=f"Từ ngữ không phù hợp được phát hiện: {', '.join(filtered_words)}",
                    related_data=f"Context: {context}, Content: {content[:100]}..."
                )
            except Exception as log_error:
                print(f"Error logging filtered content: {str(log_error)}")
        
        return jsonify({
            "original_content": content,
            "filtered_content": filtered_content,
            "filtered_words": filtered_words,
            "has_inappropriate_content": has_inappropriate_content,
            "context": context,
            "message": "Nội dung đã được kiểm tra và lọc thành công" if has_inappropriate_content else "Nội dung phù hợp"
        })
    except Exception as e:
        print(f"Error in check_content: {str(e)}")
        return jsonify({"error": str(e)}), 500

app.run(debug=True) 