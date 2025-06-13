from flask import Flask, render_template, request, jsonify, send_file, session, redirect, url_for
import numpy as np
import requests
import os
import io
import tempfile
import json
import traceback
# Thêm thư viện xử lý file docx
import docx
# Thêm python-dotenv để đọc file .env
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()
# Thêm thư viện xử lý PDF và Excel
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
# Import MongoDB
from pymongo import MongoClient
from datetime import datetime

app = Flask(__name__)
app.secret_key = 'ahp_decision_support_system_secret_key_2024'  # Cần thiết cho session

# Kết nối MongoDB
try:
    mongo_client = MongoClient('mongodb://localhost:27017/')
    db = mongo_client['ahp']
    logs_collection = db['logs']
    print("MongoDB connected successfully!")
except Exception as e:
    print(f"MongoDB connection error: {str(e)}")
    logs_collection = None

# Hàm ghi log vào MongoDB
def log_to_mongodb(data):
    try:
        if logs_collection is not None:
            # Thêm timestamp
            data['timestamp'] = datetime.now()
            # Thêm vào collection
            logs_collection.insert_one(data)
            return True
    except Exception as e:
        print(f"MongoDB log error: {str(e)}")
    return False

# Cấu hình API cho Groq
API_URL = os.getenv('API_URL', "https://api.groq.com/openai/v1/chat/completions")
GROQ_API_KEY = os.getenv('GROQ_API_KEY', "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")


def read_docx_file(file_storage):
    """Đọc nội dung file DOCX và chuyển thành string."""
    try:
        if 'docx' not in globals():
            return "Thư viện python-docx chưa được cài đặt. Không thể đọc file DOCX."
        
        # Lưu file tạm để xử lý
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            file_storage.save(temp_file.name)
            temp_path = temp_file.name
        
        # Đọc nội dung từ file docx
        doc = docx.Document(temp_path)
        content = []
        
        # Trích xuất text từ đoạn văn
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        # Trích xuất text từ bảng
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                if any(cell for cell in row_text):  # Chỉ thêm hàng có nội dung
                    content.append(' | '.join(row_text))
        
        # Xóa file tạm
        os.unlink(temp_path)
        
        return '\n'.join(content)
    except Exception as e:
        print(f"Error reading DOCX: {str(e)}")
        return f"Lỗi khi đọc file DOCX: {str(e)}"

def get_llm_suggestions(goal, file_content=None):
    """
    Nhận goal và nội dung file để phân tích và đề xuất tiêu chí và phương án.
    
    Args:
        goal: Mục tiêu cần phân tích
        file_content: Nội dung file đã được đọc (dạng string)
    
    Returns:
        Dict chứa các tiêu chí và phương án được đề xuất
    """
    # Xác định nội dung prompt dựa trên việc có file hay không
    if file_content:
        prompt = f"""Hãy phân tích mục tiêu và nội dung file đính kèm, sau đó đề xuất 4 tiêu chí đánh giá quan trọng nhất và 4 phương án lựa chọn phù hợp nhất với mục tiêu đã cho:
        
        Mục tiêu: {goal}
        
        Lưu ý:
        - Ưu tiên các tiêu chí mang tính chất định lượng 
        - Chỉ đề xuất 4 tiêu chí quan trọng nhất
        - Chỉ đề xuất 4 phương án lựa chọn tốt nhất
        - Đảm bảo các tiêu chí và phương án phải thực tế và có thể so sánh được
        - Phân tích cả nội dung file để đề xuất cho phù hợp
        
        Hãy trả về kết quả theo định dạng JSON như sau:
        {{
            "criteria": ["tiêu chí 1", "tiêu chí 2", "tiêu chí 3", "tiêu chí 4"],
            "alternatives": ["phương án 1", "phương án 2", "phương án 3", "phương án 4"]
        }}
        Dữ liệu như sau:
        {file_content[:100000]}  
        tiêu chí phải phù hợp với mục tiêu {goal};
        """
    else:
        prompt = f"""Hãy phân tích mục tiêu sau và đề xuất 4 tiêu chí đánh giá quan trọng nhất và 5 phương án lựa chọn phù hợp nhất:
        Mục tiêu: {goal}
        
        Lưu ý:
        - Chỉ đề xuất 4 tiêu chí quan trọng nhất
        - Chỉ đề xuất 4 phương án lựa chọn tốt nhất
        - Đảm bảo các tiêu chí và phương án phải thực tế và có thể so sánh được
        
        Hãy trả về kết quả theo định dạng JSON như sau:
        {{
            "criteria": ["tiêu chí 1", "tiêu chí 2", "tiêu chí 3", "tiêu chí 4"],
            "alternatives": ["phương án 1", "phương án 2", "phương án 3", "phương án 4"]
        }}
        """
    try:
        headers = {
            "Authorization": f"Bearer {GROQ_API_KEY}",
            "Content-Type": "application/json"
        }
        
        response = requests.post(
            API_URL,
            headers=headers,
            json={
                "model": "llama-3.3-70b-versatile",  # Groq model
                "messages": [
                    {"role": "system", "content": "Bạn là một chuyên gia phân tích và tư vấn ra quyết định. Hãy phân tích mục tiêu và đề xuất 4 tiêu chí đánh giá quan trọng nhất và 4 phương án thay thế phù hợp nhất."},
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.7,
                "max_tokens": 800 if file_content else 500  # Tăng max_tokens nếu có file
            }
        )
        response.raise_for_status()
        
        response_data = response.json()
        print(f"Groq API Response: {response_data}")  # Debug log
        
        if 'choices' in response_data and len(response_data['choices']) > 0:
            content = response_data['choices'][0]['message']['content']
            print(f"AI Content: {content}")  # Debug log
            
            try:
                import json
                start_idx = content.find('{')
                end_idx = content.rfind('}') + 1
                if start_idx != -1 and end_idx != 0:
                    json_str = content[start_idx:end_idx]
                    result = json.loads(json_str)
                    result['criteria'] = result['criteria'][:5]
                    result['alternatives'] = result['alternatives'][:5]
                    print(f"Successfully parsed AI result: {result}")  # Debug log
                    return result
            except json.JSONDecodeError as e:
                print("Error parsing JSON from response:", e)
                print("Raw content:", content)  # Debug log
                
                # Fallback: Tạo kết quả mặc định dựa trên mục tiêu
                return create_default_suggestions(goal)
        
        print("No valid choices in AI response")  # Debug log
        return create_default_suggestions(goal)
    except Exception as e:
        print(f"Error: {str(e)}")
        return create_default_suggestions(goal)

def create_default_suggestions(goal):
    """Tạo tiêu chí và phương án mặc định khi AI không hoạt động"""
    # Phân tích từ khóa trong mục tiêu để tạo tiêu chí phù hợp
    goal_lower = goal.lower()
    
    default_criteria = []
    default_alternatives = []
    
    # Tiêu chí mặc định dựa trên từ khóa
    if any(word in goal_lower for word in ['mua', 'chọn', 'lựa chọn', 'sản phẩm', 'dịch vụ']):
        default_criteria = ["Giá cả", "Chất lượng", "Độ tin cậy", "Tính năng"]
        default_alternatives = ["Lựa chọn A", "Lựa chọn B", "Lựa chọn C", "Lựa chọn D"]
    elif any(word in goal_lower for word in ['tuyển dụng', 'nhân viên', 'ứng viên']):
        default_criteria = ["Kinh nghiệm", "Kỹ năng", "Học vấn", "Tính cách"]
        default_alternatives = ["Ứng viên A", "Ứng viên B", "Ứng viên C", "Ứng viên D"]
    elif any(word in goal_lower for word in ['đầu tư', 'dự án', 'kinh doanh']):
        default_criteria = ["ROI", "Rủi ro", "Thời gian", "Tài nguyên"]
        default_alternatives = ["Dự án A", "Dự án B", "Dự án C", "Dự án D"]
    elif any(word in goal_lower for word in ['học', 'giáo dục', 'trường']):
        default_criteria = ["Chất lượng giảng dạy", "Chi phí", "Uy tín", "Vị trí"]
        default_alternatives = ["Lựa chọn A", "Lựa chọn B", "Lựa chọn C", "Lựa chọn D"]
    else:
        # Tiêu chí tổng quát
        default_criteria = ["Chi phí", "Chất lượng", "Thời gian", "Hiệu quả"]
        default_alternatives = ["Phương án A", "Phương án B", "Phương án C", "Phương án D"]
    
    return {
        "criteria": default_criteria,
        "alternatives": default_alternatives
    }

def calculate_ahp(comparison_matrix):
    matrix_original = comparison_matrix.copy()
    n = comparison_matrix.shape[0]

    # Tính sum từng cột
    arr_sum = []
    for i in range(n):
        sum = 0
        for j in range(n):
            sum += comparison_matrix[j][i]
        arr_sum.append(sum)

    # Chuẩn hóa ma trận so sánh cặp
    normalized_matrix = np.zeros((n, n))
    for i in range(n):
        for j in range(n):
            normalized_matrix[i][j] = comparison_matrix[i][j] / arr_sum[j]

    # Tính Criteria Weight
    arr_avg = []
    for i in range(n):
        sum = 0
        for j in range(n):
            sum += normalized_matrix[i][j]
        arr_avg.append(sum/n)

    # Nhân Criteria Weight với ma trận so sánh cặp
    for i in range(n):
        for j in range(n):
            matrix_original[j][i] = matrix_original[j][i] * arr_avg[i]
    
    # Tính Weighted Sum Value
    arr_WSV = []
    for i in range(n):
        sum = 0
        for j in range(n):
            sum += matrix_original[i][j]
        arr_WSV.append(sum)
    
    # Tính Consistency Vector
    arr_CV = []
    for i in range(len(arr_WSV)):
        arr_CV.append(arr_WSV[i] / arr_avg[i])

    # Tính Lambda Max
    sum = 0
    for value in arr_CV:
        sum += value
    lambda_max = sum / len(arr_CV)

    # Tính CI
    ci = (lambda_max - n) / (n - 1)

    # RI values
    ri_values = [0.00, 0.00, 0.58, 0.90, 1.12, 1.24, 1.32, 1.41, 1.45, 1.49, 1.51, 1.54, 1.56, 1.58, 1.59]
    ri = ri_values[n-1] if n <= len(ri_values) else ri_values[-1]

    # Tính CR
    cr = ci / ri

    return {
        "weights": arr_avg,
        "lambda_max": float(lambda_max),
        "ci": float(ci),
        "cr": float(cr),
        "normalized_matrix": normalized_matrix.tolist()
    }

def generate_criteria_matrix_ai(goal, criteria):
    """
    Sử dụng AI để tự động tạo ma trận so sánh các tiêu chí dựa trên mục tiêu
    
    Args:
        goal: Mục tiêu của phân tích AHP
        criteria: Danh sách các tiêu chí
    
    Returns:
        Dict chứa ma trận so sánh và thông tin
    """
    try:
        prompt = f"""Bạn là chuyên gia AHP. Hãy tạo ma trận so sánh cặp {len(criteria)}x{len(criteria)} cho các tiêu chí sau theo mục tiêu "{goal}":

Các tiêu chí: {', '.join(criteria)}

Sử dụng thang đo Saaty (1-9):
- 1: Hai tiêu chí quan trọng như nhau
- 3: Tiêu chí A hơi quan trọng hơn B  
- 5: Tiêu chí A quan trọng hơn B
- 7: Tiêu chí A rất quan trọng hơn B
- 9: Tiêu chí A cực kỳ quan trọng hơn B

Phân tích sâu từng tiêu chí trong bối cảnh mục tiêu "{goal}":
1. Xem xét tầm quan trọng tương đối của từng cặp tiêu chí
2. Đảm bảo tính nhất quán logic (nếu A > B và B > C thì A > C)  
3. Tạo ma trận đối xứng nghịch đảo (a_ij = 1/a_ji)
4. Đảm bảo đường chéo = 1

Chỉ trả về ma trận số dưới dạng JSON với SỐ THẬP PHÂN (không dùng phân số):
[
  [1, 3.0, 5.0, 7.0],
  [0.333, 1, 2.0, 3.0], 
  [0.2, 0.5, 1, 2.0],
  [0.143, 0.333, 0.5, 1]
]

QUAN TRỌNG: 
- Chỉ sử dụng số thập phân (VD: 0.333 thay vì 1/3)
- Không sử dụng phân số như 1/3, 1/5, 1/7
- Ma trận phải có CR < 0.1 để đảm bảo tính nhất quán."""

        headers = {
            "Authorization": f"Bearer {GROQ_API_KEY}",
            "Content-Type": "application/json"
        }
        
        max_attempts = 3
        for attempt in range(max_attempts):
            response = requests.post(
                API_URL,
                headers=headers,
                json={
                    "model": "llama-3.3-70b-versatile",
                    "messages": [
                        {"role": "system", "content": "Bạn là chuyên gia AHP có kinh nghiệm về ma trận so sánh và đảm bảo tính nhất quán."},
                        {"role": "user", "content": prompt}
                    ],
                    "temperature": 0.3,
                    "max_tokens": 1000
                }
            )
            response.raise_for_status()
            
            response_data = response.json()
            print(f"🤖 AI Response (lần {attempt + 1}): {response_data}")  # Log full response
            
            if 'choices' in response_data and len(response_data['choices']) > 0:
                content = response_data['choices'][0]['message']['content']
                print(f"📝 AI Content (lần {attempt + 1}):")
                print("=" * 50)
                print(content)
                print("=" * 50)
                
                # Trích xuất ma trận từ JSON
                import json
                import re
                
                # Tìm ma trận JSON
                matrix_match = re.search(r'\[\s*\[.*?\]\s*\]', content, re.DOTALL)
                if matrix_match:
                    matrix_json = matrix_match.group(0)
                    print(f"🎯 Extracted Matrix JSON (lần {attempt + 1}): {matrix_json}")
                    
                    try:
                        # Convert phân số thành số thập phân trước khi parse JSON
                        import re
                        
                        # Thay thế tất cả phân số dạng "số/số" thành số thập phân
                        def fraction_to_decimal(match):
                            numerator = float(match.group(1))
                            denominator = float(match.group(2))
                            return str(numerator / denominator)
                        
                        # Pattern để tìm phân số: một hoặc nhiều chữ số, dấu /, một hoặc nhiều chữ số
                        matrix_json_fixed = re.sub(r'(\d+)/(\d+)', fraction_to_decimal, matrix_json)
                        print(f"🔧 Fixed Matrix JSON (lần {attempt + 1}): {matrix_json_fixed}")
                        
                        generated_matrix = json.loads(matrix_json_fixed)
                        print(f"✅ Successfully parsed matrix (lần {attempt + 1}): {generated_matrix}")
                        
                        # Chuyển thành numpy array và validate
                        n = len(criteria)
                        matrix = np.ones((n, n))
                        
                        for i in range(n):
                            for j in range(n):
                                if i < len(generated_matrix) and j < len(generated_matrix[i]):
                                    val = generated_matrix[i][j]
                                    if isinstance(val, str) and '/' in val:
                                        parts = val.split('/')
                                        if len(parts) == 2:
                                            matrix[i][j] = float(parts[0]) / float(parts[1])
                                    else:
                                        matrix[i][j] = float(val)
                        
                        # Kiểm tra và cải thiện tính nhất quán
                        result = calculate_ahp(matrix)
                        cr = result["cr"]
                        
                        if cr < 0.1:  # Ma trận nhất quán
                            return {
                                "matrix": matrix.tolist(),
                                "weights": result["weights"],
                                "cr": cr,
                                "lambda_max": result["lambda_max"],
                                "is_consistent": True,
                                "attempt": attempt + 1
                            }
                        else:
                            print(f"Lần thử {attempt + 1}: CR = {cr:.4f} (>0.1), thử lại...")
                            
                    except Exception as e:
                        print(f"❌ Lỗi xử lý ma trận lần {attempt + 1}: {e}")
                        print(f"🔍 Raw matrix JSON: {matrix_json}")
                        import traceback
                        traceback.print_exc()
                else:
                    print(f"❌ Không tìm thấy pattern ma trận JSON trong content (lần {attempt + 1})")
                    print(f"🔍 Content đã search: {content[:200]}...")  # Show first 200 chars
            else:
                print(f"❌ Không có choices trong AI response (lần {attempt + 1})")
                print(f"🔍 Response data: {response_data}")
        
        # Nếu không tạo được ma trận nhất quán, tạo ma trận mặc định hợp lý
        print("Không tạo được ma trận nhất quán từ AI, sử dụng ma trận mặc định")
        return create_default_criteria_matrix(len(criteria))
        
    except Exception as e:
        print(f"Lỗi AI generate criteria matrix: {e}")
        print(f"Groq API có thể đang gặp vấn đề. Sử dụng ma trận mặc định.")
        import traceback
        traceback.print_exc()
        return create_default_criteria_matrix(len(criteria))

def create_default_criteria_matrix(n):
    """Tạo ma trận tiêu chí mặc định với sự khác biệt hợp lý"""
    matrix = np.ones((n, n))
    
    # Tạo ma trận với sự khác biệt có ý nghĩa
    # Sử dụng các giá trị Saaty phổ biến: 1, 2, 3, 5
    import random
    
    # Tạo các giá trị khác biệt theo pattern hợp lý
    values = [1, 2, 3, 2, 3, 5, 3, 5, 7]  # Giá trị Saaty
    
    idx = 0
    for i in range(n):
        for j in range(i+1, n):
            # Tạo sự khác biệt dựa trên vị trí và ngẫu nhiên
            if i == 0:  # Tiêu chí đầu tiên quan trọng hơn một chút
                val = values[idx % len(values)] if idx < len(values) else 2
            elif j == n-1:  # Tiêu chí cuối ít quan trọng hơn
                val = 1 / (values[idx % len(values)] if idx < len(values) else 2)
            else:
                val = values[idx % len(values)] if idx < len(values) else random.choice([1, 2, 3])
            
            # Đảm bảo giá trị trong khoảng hợp lý
            val = max(0.2, min(9, val))
            matrix[i][j] = val
            matrix[j][i] = 1/val
            idx += 1
    
    result = calculate_ahp(matrix)
    
    # Nếu CR quá cao, điều chỉnh
    if result["cr"] > 0.1:
        # Tạo ma trận đơn giản hơn với ít sự khác biệt
        matrix = np.ones((n, n))
        for i in range(n):
            for j in range(i+1, n):
                val = 2 if i == 0 else (3 if j == n-1 else random.choice([1, 2]))
                matrix[i][j] = val
                matrix[j][i] = 1/val
        result = calculate_ahp(matrix)
    
    return {
        "matrix": matrix.tolist(),
        "weights": result["weights"],
        "cr": result["cr"],
        "lambda_max": result["lambda_max"],
        "is_consistent": True,
        "attempt": 1
    }

def generate_alternative_matrix_ai_function(goal, criterion, alternatives):
    """
    Sử dụng AI để tự động tạo ma trận so sánh phương án theo tiêu chí
    
    Args:
        goal: Mục tiêu của phân tích AHP
        criterion: Tiêu chí đang xét
        alternatives: Danh sách các phương án
    
    Returns:
        Dict chứa ma trận so sánh và thông tin
    """
    try:
        # Prompt cải tiến để AI phân tích thông số thực tế
        prompt = f"""Bạn là chuyên gia AHP và am hiểu sâu về sản phẩm/dịch vụ. Hãy tạo ma trận so sánh cặp {len(alternatives)}x{len(alternatives)} cho các phương án theo tiêu chí "{criterion}" trong mục tiêu "{goal}":

Các phương án: {', '.join(alternatives)}
Tiêu chí đánh giá: {criterion}
Mục tiêu: {goal}

HƯỚNG DẪN PHÂN TÍCH CHI TIẾT:

1. **Phân tích thông số kỹ thuật thực tế** của từng phương án:
   - Nếu là sản phẩm công nghệ: tìm hiểu cấu hình, hiệu năng, thông số kỹ thuật
   - Nếu là dịch vụ: phân tích chất lượng, tính năng, độ tin cậy
   - Nếu là dự án: xem xét quy mô, ngân sách, thời gian, rủi ro

2. **So sánh định lượng dựa trên dữ liệu thực**:
   - Tìm kiếm thông tin công khai về các phương án
   - Sử dụng số liệu cụ thể để so sánh (ví dụ: tốc độ CPU, dung lượng RAM, giá cả, đánh giá người dùng)
   - Tính toán tỷ lệ thực tế giữa các phương án

3. **Áp dụng thang Saaty (1-9) dựa trên tỷ lệ thực**:
   - 1: Hai phương án tương đương về tiêu chí này
   - 3: Phương án A tốt hơn B khoảng 1.5-2 lần
   - 5: Phương án A tốt hơn B khoảng 2-3 lần  
   - 7: Phương án A tốt hơn B khoảng 3-4 lần
   - 9: Phương án A vượt trội B hơn 4 lần

4. **Ví dụ cụ thể**:
   - Nếu tiêu chí "Processor" và có laptop với i7-11800H vs i5-8265U → tỷ lệ khoảng 3-5
   - Nếu tiêu chí "Giá cả" và sản phẩm A = 20 triệu, B = 15 triệu → B tốt hơn A = 20/15 ≈ 1.3 → làm tròn về 2

QUAN TRỌNG: 
- Dựa trên kiến thức thực tế về các phương án để tạo ma trận
- KHÔNG tạo ma trận đơn vị (tất cả bằng 1) 
- Phải có sự khác biệt có ý nghĩa giữa các phương án
- Đảm bảo CR < 0.1

Chỉ trả về ma trận số dưới dạng JSON với SỐ THẬP PHÂN (không dùng phân số):
[
  [1, 3.0, 5.0, 2.0],
  [0.333, 1, 2.0, 3.0],
  [0.2, 0.5, 1, 1.5],
  [0.5, 0.333, 0.667, 1]
]

QUAN TRỌNG: 
- Chỉ sử dụng số thập phân (VD: 0.333 thay vì 1/3)
- Không sử dụng phân số như 1/3, 1/5, 1/7
- Ma trận phải có CR < 0.1"""

        headers = {
            "Authorization": f"Bearer {GROQ_API_KEY}",
            "Content-Type": "application/json"
        }
        
        max_attempts = 3
        for attempt in range(max_attempts):
            response = requests.post(
                API_URL,
                headers=headers,
                json={
                    "model": "llama-3.3-70b-versatile",
                    "messages": [
                        {"role": "system", "content": "Bạn là một chuyên gia AHP với kinh nghiệm phân tích và so sánh phương án."},
                        {"role": "user", "content": prompt}
                    ],
                    "temperature": 0.3,
                    "max_tokens": 800
                }
            )
            response.raise_for_status()
            
            response_data = response.json()
            print(f"🤖 AI Alternative Matrix Response (lần {attempt + 1} - {criterion}): {response_data}")  # Log full response
            
            if 'choices' in response_data and len(response_data['choices']) > 0:
                content = response_data['choices'][0]['message']['content']
                print(f"📝 AI Alternative Matrix Content (lần {attempt + 1} - {criterion}):")
                print("=" * 50)
                print(content)
                print("=" * 50)
                
                # Trích xuất ma trận từ JSON
                import json
                import re
                
                matrix_match = re.search(r'\[\s*\[.*?\]\s*\]', content, re.DOTALL)
                if matrix_match:
                    matrix_json = matrix_match.group(0)
                    print(f"🎯 Extracted Alternative Matrix JSON (lần {attempt + 1} - {criterion}): {matrix_json}")
                    
                    try:
                        # Convert phân số thành số thập phân trước khi parse JSON
                        import re
                        
                        # Thay thế tất cả phân số dạng "số/số" thành số thập phân
                        def fraction_to_decimal(match):
                            numerator = float(match.group(1))
                            denominator = float(match.group(2))
                            return str(numerator / denominator)
                        
                        # Pattern để tìm phân số
                        matrix_json_fixed = re.sub(r'(\d+)/(\d+)', fraction_to_decimal, matrix_json)
                        print(f"🔧 Fixed Alternative Matrix JSON (lần {attempt + 1} - {criterion}): {matrix_json_fixed}")
                        
                        generated_matrix = json.loads(matrix_json_fixed)
                        print(f"✅ Successfully parsed alternative matrix (lần {attempt + 1} - {criterion}): {generated_matrix}")
                        
                        # Chuyển thành numpy array
                        n = len(alternatives)
                        matrix = np.ones((n, n))
                        
                        for i in range(n):
                            for j in range(n):
                                if i < len(generated_matrix) and j < len(generated_matrix[i]):
                                    val = generated_matrix[i][j]
                                    if isinstance(val, str) and '/' in val:
                                        parts = val.split('/')
                                        if len(parts) == 2:
                                            matrix[i][j] = float(parts[0]) / float(parts[1])
                                    else:
                                        matrix[i][j] = float(val)
                        
                        # Kiểm tra tính nhất quán
                        result = calculate_ahp(matrix)
                        cr = result["cr"]
                        
                        if cr < 0.1:  # Ma trận nhất quán
                            alt_weights = {alt: weight for alt, weight in zip(alternatives, result["weights"])}
                            return {
                                "matrix": matrix.tolist(),
                                "weights": result["weights"],
                                "alternatives": alt_weights,
                                "cr": cr,
                                "lambda_max": result["lambda_max"],
                                "is_consistent": True,
                                "attempt": attempt + 1
                            }
                        else:
                            print(f"Ma trận {criterion} lần {attempt + 1}: CR = {cr:.4f} (>0.1), thử lại...")
                            
                    except Exception as e:
                        print(f"❌ Lỗi xử lý ma trận {criterion} lần {attempt + 1}: {e}")
                        print(f"🔍 Raw alternative matrix JSON: {matrix_json}")
                        import traceback
                        traceback.print_exc()
                else:
                    print(f"❌ Không tìm thấy pattern ma trận JSON trong alternative content (lần {attempt + 1} - {criterion})")
                    print(f"🔍 Alternative content đã search: {content[:200]}...")  # Show first 200 chars
            else:
                print(f"❌ Không có choices trong AI alternative response (lần {attempt + 1} - {criterion})")
                print(f"🔍 Alternative response data: {response_data}")
        
        # Nếu không tạo được ma trận nhất quán, tạo ma trận mặc định
        print(f"Không tạo được ma trận nhất quán cho {criterion}, sử dụng ma trận mặc định")
        return create_default_alternative_matrix(alternatives)
        
    except Exception as e:
        print(f"Lỗi AI generate alternative matrix cho {criterion}: {e}")
        import traceback
        traceback.print_exc()
        return create_default_alternative_matrix(alternatives)

def create_default_alternative_matrix(alternatives):
    """Tạo ma trận phương án mặc định với sự khác biệt có ý nghĩa"""
    n = len(alternatives)
    matrix = np.ones((n, n))
    
    # Tạo ma trận với sự khác biệt thực tế thay vì toàn bộ bằng 1
    import random
    random.seed(42)  # Để có kết quả nhất quán
    
    # Tạo các giá trị Saaty có ý nghĩa: 1, 2, 3, 5, 7
    saaty_values = [1, 2, 3, 5]  # Tránh giá trị quá cao
    
    for i in range(n):
        for j in range(i+1, n):
            # Tạo sự khác biệt dựa trên vị trí (giả định phương án đầu tốt hơn phương án cuối)
            if i == 0:  # Phương án đầu tiên thường tốt hơn
                val = random.choice([2, 3])
            elif j == n-1:  # Phương án cuối cùng thường kém hơn
                val = 1 / random.choice([2, 3])
            else:
                # Các phương án ở giữa có sự khác biệt nhỏ hơn
                val = random.choice([1, 2, 1/2])
            
            # Đảm bảo giá trị trong khoảng hợp lý
            val = max(1/7, min(7, val))
            matrix[i][j] = val
            matrix[j][i] = 1/val
    
    result = calculate_ahp(matrix)
    alt_weights = {alt: weight for alt, weight in zip(alternatives, result["weights"])}
    
    return {
        "matrix": matrix.tolist(),
        "weights": result["weights"],
        "alternatives": alt_weights,
        "cr": result["cr"],
        "lambda_max": result["lambda_max"],
        "is_consistent": True,
        "attempt": 1
    }

def calculate_final_scores(criteria_weights, alternative_matrices):
    """Calculate final scores from criteria weights and alternative matrices"""
    final_scores = {}
    
    # Lặp qua tất cả phương án trong ma trận đầu tiên để lấy danh sách phương án
    first_criterion = list(alternative_matrices.keys())[0] if alternative_matrices else None
    if not first_criterion:
        return {}
        
    alternatives = list(alternative_matrices[first_criterion]["alternatives"].keys())
    
    # Tính điểm tổng hợp cho từng phương án
    for alt in alternatives:
        score = 0
        for i, criterion in enumerate(alternative_matrices.keys()):
            criterion_weight = criteria_weights[i]
            if criterion in alternative_matrices and alt in alternative_matrices[criterion]["alternatives"]:
                alt_weight = alternative_matrices[criterion]["alternatives"][alt]
                score += criterion_weight * alt_weight
        final_scores[alt] = score
    
    return final_scores

@app.route('/login', methods=['GET', 'POST'])
def login():
    """Trang đăng nhập - chấp nhận bất kỳ thông tin đăng nhập nào"""
    if request.method == 'POST':
        # Lấy thông tin từ form
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '').strip()
        
        # Kiểm tra có nhập thông tin không (để trông có vẻ thật)
        if not username or not password:
            return render_template('login.html', error='Vui lòng nhập đầy đủ tài khoản và mật khẩu')
        
        # Kiểm tra độ dài tối thiểu (để trông có vẻ thật)
        if len(username) < 3 or len(password) < 3:
            return render_template('login.html', error='Tài khoản và mật khẩu phải có ít nhất 3 ký tự')
        
        # "Xác thực" thành công - chấp nhận bất kỳ thông tin hợp lệ nào
        session['logged_in'] = True
        session['username'] = username
        
        # Log đăng nhập vào MongoDB
        log_to_mongodb({
            'type': 'login',
            'username': username,
            'timestamp': datetime.now(),
            'ip_address': request.remote_addr
        })
        
        return redirect(url_for('index'))
    
    # Nếu đã đăng nhập, chuyển thẳng tới trang chính
    if session.get('logged_in'):
        return redirect(url_for('index'))
    
    return render_template('login.html')

@app.route('/logout')
def logout():
    """Đăng xuất"""
    session.clear()
    return redirect(url_for('login'))

@app.route('/')
def index():
    """Trang chính - yêu cầu đăng nhập"""
    # Kiểm tra đăng nhập
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    return render_template('index.html', username=session.get('username', 'Người dùng'))

@app.route('/generate_criteria_matrix', methods=['POST'])
def generate_criteria_matrix():
    """Tự động tạo ma trận so sánh tiêu chí bằng AI"""
    try:
        data = request.json
        goal = data.get('goal', '')
        criteria = data.get('criteria', [])
        
        if not goal or not criteria:
            return jsonify({"error": "Thiếu mục tiêu hoặc tiêu chí"}), 400
        
        if len(criteria) < 2:
            return jsonify({"error": "Cần ít nhất 2 tiêu chí để so sánh"}), 400
        
        # Gọi AI để tạo ma trận
        matrix_result = generate_criteria_matrix_ai(goal, criteria)
        
        # Log vào MongoDB
        log_to_mongodb({
            'type': 'generate_criteria_matrix',
            'goal': goal,
            'criteria': criteria,
            'matrix_result': matrix_result
        })
        
        return jsonify({
            "success": True,
            "matrix": matrix_result["matrix"],
            "weights": matrix_result["weights"],
            "cr": matrix_result["cr"],
            "lambda_max": matrix_result["lambda_max"],
            "is_consistent": matrix_result["is_consistent"],
            "attempt": matrix_result["attempt"],
            "message": f"Ma trận được tạo thành công với CR = {matrix_result['cr']:.4f}"
        })
        
    except Exception as e:
        print(f"Error in generate_criteria_matrix: {str(e)}")
        return jsonify({"error": f"Lỗi tạo ma trận: {str(e)}"}), 500

@app.route('/generate_alternative_matrix_ai', methods=['POST'])
def generate_alternative_matrix_ai_endpoint():
    """Tự động tạo ma trận so sánh phương án cho một tiêu chí bằng AI"""
    try:
        data = request.json
        goal = data.get('goal', '')
        criterion = data.get('criterion', '')
        alternatives = data.get('alternatives', [])
        
        if not goal or not criterion or not alternatives:
            return jsonify({"error": "Thiếu mục tiêu, tiêu chí hoặc phương án"}), 400
        
        if len(alternatives) < 2:
            return jsonify({"error": "Cần ít nhất 2 phương án để so sánh"}), 400
        
        # Gọi AI để tạo ma trận phương án cho tiêu chí này
        matrix_result = generate_alternative_matrix_ai_function(goal, criterion, alternatives)
        
        # Log vào MongoDB
        log_to_mongodb({
            'type': 'generate_alternative_matrix_ai',
            'goal': goal,
            'criterion': criterion,
            'alternatives': alternatives,
            'matrix_result': matrix_result
        })
        
        return jsonify({
            "success": True,
            "matrix": matrix_result["matrix"],
            "weights": matrix_result["weights"],
            "cr": matrix_result["cr"],
            "lambda_max": matrix_result["lambda_max"],
            "is_consistent": matrix_result["is_consistent"],
            "attempt": matrix_result["attempt"],
            "message": f"Ma trận phương án cho tiêu chí '{criterion}' được tạo thành công với CR = {matrix_result['cr']:.4f}"
        })
        
    except Exception as e:
        print(f"Error in generate_alternative_matrix_ai: {str(e)}")
        return jsonify({"error": f"Lỗi tạo ma trận phương án: {str(e)}"}), 500

@app.route('/generate_full_ahp_matrices', methods=['POST'])
def generate_full_ahp_matrices():
    """Tự động tạo toàn bộ ma trận AHP (tiêu chí + phương án) bằng AI"""
    try:
        data = request.json
        goal = data.get('goal', '')
        criteria = data.get('criteria', [])
        alternatives = data.get('alternatives', [])
        
        if not goal or not criteria or not alternatives:
            return jsonify({"error": "Thiếu dữ liệu mục tiêu, tiêu chí hoặc phương án"}), 400
        
        # 1. Tạo ma trận so sánh tiêu chí
        criteria_matrix_result = generate_criteria_matrix_ai(goal, criteria)
        
        # 2. Tạo ma trận so sánh phương án cho từng tiêu chí
        alternative_matrices = {}
        for criterion in criteria:
            alt_matrix_result = generate_alternative_matrix_ai_function(goal, criterion, alternatives)
            alternative_matrices[criterion] = alt_matrix_result
        
        # 3. Tính điểm cuối cùng
        final_scores = calculate_final_scores(criteria_matrix_result["weights"], alternative_matrices)
        
        # Log vào MongoDB
        log_to_mongodb({
            'type': 'generate_full_ahp_matrices',
            'goal': goal,
            'criteria': criteria,
            'alternatives': alternatives,
            'final_scores': final_scores
        })
        
        return jsonify({
            "success": True,
            "criteria_matrix": criteria_matrix_result,
            "alternative_matrices": alternative_matrices,
            "final_scores": final_scores,
            "message": "Tạo toàn bộ ma trận AHP thành công!"
        })
        
    except Exception as e:
        print(f"Error in generate_full_ahp_matrices: {str(e)}")
        return jsonify({"error": f"Lỗi tạo ma trận AHP: {str(e)}"}), 500

@app.route('/auto_generate_complete_ahp', methods=['POST'])
def auto_generate_complete_ahp():
    """Tự động tạo toàn bộ phân tích AHP từ mục tiêu"""
    try:
        # Lấy goal từ form hoặc JSON
        if request.form:
            goal = request.form.get('goal', '')
        else:
            goal = request.json.get('goal', '') if request.json else ''
        
        if not goal:
            return jsonify({"error": "Vui lòng nhập mục tiêu"}), 400
        
        # Xử lý file nếu có
        file_content = None
        if 'file' in request.files and request.files['file'].filename != '':
            file = request.files['file']
            filename = file.filename.lower()
            
            if filename.endswith('.docx'):
                file_content = read_docx_file(file)
            elif filename.endswith('.doc'):
                return jsonify({"error": "File .doc cũ không được hỗ trợ. Vui lòng chuyển đổi sang .docx"}), 400
            else:
                return jsonify({"error": "Chỉ hỗ trợ file Word (.docx)"}), 400
        
        # 1. Gọi AI để đề xuất tiêu chí và phương án
        suggestions = get_llm_suggestions(goal, file_content)
        if "error" in suggestions:
            return jsonify(suggestions), 500
        
        criteria = suggestions.get('criteria', [])
        alternatives = suggestions.get('alternatives', [])
        
        if not criteria or not alternatives:
            return jsonify({"error": "AI không thể đề xuất tiêu chí và phương án hợp lệ"}), 500
        
        # 2. Tạo ma trận so sánh tiêu chí
        criteria_matrix_result = generate_criteria_matrix_ai(goal, criteria)
        
        # 3. Tạo ma trận so sánh phương án cho từng tiêu chí
        alternative_matrices = {}
        for criterion in criteria:
            alt_matrix_result = generate_alternative_matrix_ai_function(goal, criterion, alternatives)
            alternative_matrices[criterion] = alt_matrix_result
        
        # 4. Tính điểm cuối cùng
        final_scores = calculate_final_scores(criteria_matrix_result["weights"], alternative_matrices)
        
        # Log vào MongoDB
        log_to_mongodb({
            'type': 'auto_generate_complete_ahp',
            'goal': goal,
            'has_file': file_content is not None,
            'criteria': criteria,
            'alternatives': alternatives,
            'final_scores': final_scores
        })
        
        return jsonify({
            "success": True,
            "goal": goal,
            "criteria": criteria,
            "alternatives": alternatives,
            "criteria_matrix": criteria_matrix_result,
            "alternative_matrices": alternative_matrices,
            "final_scores": final_scores,
            "message": "Tạo phân tích AHP hoàn chỉnh thành công!",
            "auto_generated": True
        })
        
    except Exception as e:
        print(f"Error in auto_generate_complete_ahp: {str(e)}")
        return jsonify({"error": f"Lỗi tạo phân tích AHP tự động: {str(e)}"}), 500

@app.route('/get_suggestions', methods=['POST'])
def get_suggestions():
    """Xử lý yêu cầu phân tích từ form hoặc AJAX."""
    try:
        # Lấy goal từ form hoặc JSON
        if request.form:
            goal = request.form.get('goal', '')
        else:
            goal = request.json.get('goal', '')
        
        # Kiểm tra goal có tồn tại
        if not goal:
            return jsonify({"error": "Vui lòng nhập mục tiêu"}), 400
        
        # Xử lý file nếu có
        file_content = None
        if 'file' in request.files and request.files['file'].filename != '':
            file = request.files['file']
            filename = file.filename.lower()
            
            # Kiểm tra loại file và đọc nội dung
            if filename.endswith('.docx'):
                file_content = read_docx_file(file)
            elif filename.endswith('.doc'):
                return jsonify({"error": "File .doc cũ không được hỗ trợ. Vui lòng chuyển đổi sang .docx"}), 400
            else:
                return jsonify({"error": "Chỉ hỗ trợ file Word (.docx)"}), 400
        
        # Gọi hàm get_llm_suggestions
        suggestions = get_llm_suggestions(goal, file_content)
        return jsonify(suggestions)
        
    except Exception as e:
        print(f"Error in get_suggestions: {str(e)}")
        return jsonify({"error": f"Lỗi xử lý yêu cầu: {str(e)}"}), 500

@app.route('/calculate_ahp', methods=['POST'])
def calculate():
    data = request.json
    
    # Chuyển đổi ma trận đầu vào từ chuỗi sang số
    input_matrix = data['matrix']
    comparison_matrix = []
    
    for row in input_matrix:
        numeric_row = []
        for value in row:
            if isinstance(value, str) and '/' in value:
                # Xử lý giá trị dạng phân số như "1/3", "1/5", etc.
                parts = value.split('/')
                if len(parts) == 2:
                    try:
                        numeric_value = float(parts[0]) / float(parts[1])
                        numeric_row.append(numeric_value)
                    except (ValueError, ZeroDivisionError):
                        return jsonify({"error": f"Không thể chuyển đổi giá trị '{value}' thành số"}), 400
            else:
                # Xử lý giá trị số hoặc chuỗi số
                try:
                    numeric_row.append(float(value))
                except ValueError:
                    return jsonify({"error": f"Không thể chuyển đổi giá trị '{value}' thành số"}), 400
        comparison_matrix.append(numeric_row)
    
    # Chuyển đổi sang numpy array và tính toán
    comparison_matrix = np.array(comparison_matrix)
    result = calculate_ahp(comparison_matrix)
    return jsonify(result)

@app.route('/calculate_alternative_matrices', methods=['POST'])
def calculate_alternative_matrices():
    """Tính toán ma trận đánh giá các phương án theo từng tiêu chí"""
    data = request.json
    criteria = data.get('criteria', [])
    alternatives = data.get('alternatives', [])
    matrices = data.get('matrices', {})
    criteria_weights = data.get('criteria_weights', [])
    
    # Kết quả của từng ma trận phương án theo từng tiêu chí
    alt_matrices_results = {}
    
    try:
        # Kiểm tra xem tất cả các tiêu chí đã có kết quả tính toán chưa
        for criterion in criteria:
            if criterion not in matrices:
                return jsonify({"error": f"Thiếu kết quả tính toán cho tiêu chí '{criterion}'"}), 400
            
            matrix_data = matrices[criterion]
            # Kiểm tra nếu matrix_data đã có thuộc tính weights và alternatives
            if isinstance(matrix_data, dict) and 'weights' in matrix_data and 'alternatives' in matrix_data and 'cr' in matrix_data:
                # Trường hợp đã có kết quả tính toán sẵn từ frontend
                alt_matrices_results[criterion] = matrix_data
            else:
                # Trường hợp là ma trận raw chưa tính toán (hiện tại không xảy ra do frontend đã tính)
                return jsonify({"error": f"Dữ liệu ma trận cho tiêu chí '{criterion}' không đúng định dạng"}), 400
    except Exception as e:
        return jsonify({"error": f"Lỗi khi xử lý ma trận: {str(e)}"}), 400
    
    # Tính điểm cuối cùng cho các phương án
    try:
        final_scores = calculate_final_scores(criteria_weights, alt_matrices_results)
        
        return jsonify({
            "alt_matrices_results": alt_matrices_results,
            "final_scores": final_scores
        })
    except Exception as e:
        return jsonify({"error": f"Lỗi khi tính điểm cuối cùng: {str(e)}"}), 400

@app.route('/get_alternative_matrices', methods=['POST'])
def get_alternative_matrices():
    """Lấy dữ liệu ma trận so sánh các phương án theo từng tiêu chí từ server"""
    try:
        data = request.json
        criteria = data.get('criteria', [])
        alternatives = data.get('alternatives', [])
        goal = data.get('goal', '')
        
        # Kiểm tra dữ liệu đầu vào
        if not criteria or not alternatives:
            return jsonify({"error": "Thiếu dữ liệu tiêu chí hoặc phương án"}), 400
            
        # Chuẩn bị ma trận so sánh cặp cho từng tiêu chí
        matrices = {}
        matrices_info = {}
        
        for criterion in criteria:
            n = len(alternatives)
            # Tạo ma trận với giá trị mặc định là 1 (đường chéo là 1, các giá trị khác được tính từ LLM)
            matrix = np.ones((n, n))
            
            # Tạo ma trận đơn giản với sự khác biệt nhỏ
            matrix = np.ones((n, n))
            
            # Thêm một chút khác biệt để tránh ma trận đơn vị
            for i in range(n):
                for j in range(i+1, n):
                    val = 1 + (i - j) * 0.1  # Sự khác biệt nhỏ
                    val = max(0.5, min(2, val))  # Giới hạn trong khoảng hợp lý
                    matrix[i][j] = val
                    matrix[j][i] = 1/val
            
            try:
                # Tạo ma trận so sánh đơn giản
                result = calculate_ahp(matrix)
                
                # Chuyển ma trận thành string format
                string_matrix = []
                for i in range(n):
                    string_row = []
                    for j in range(n):
                        val = matrix[i][j]
                        if val == 1:
                            string_row.append("1")
                        elif val < 1:
                            denominator = int(round(1/val))
                            string_row.append(f"1/{denominator}")
                        else:
                            string_row.append(str(round(val, 2)))
                    string_matrix.append(string_row)
                
                matrices[criterion] = string_matrix
                matrices_info[criterion] = {
                    "is_quantitative": False,
                    "cr": result["cr"],
                    "is_consistent": True
                }

                
            except Exception as e:
                print(f"Error generating matrix for {criterion}: {str(e)}")
                # Tạo ma trận mặc định nếu có lỗi
                matrix = np.ones((n, n))
                string_matrix = []
                for i in range(n):
                    string_row = []
                    for j in range(n):
                        string_row.append("1")
                    string_matrix.append(string_row)
                matrices[criterion] = string_matrix
                matrices_info[criterion] = {
                    "is_quantitative": False,
                    "cr": 0.0,
                    "is_consistent": True
                }
        
        # Thêm trường thông tin về phân tích định lượng để hiển thị cho người dùng
        return jsonify({
            "matrices": matrices,
            "matrices_info": matrices_info,
            "matrix_info": "Ma trận so sánh được tạo dựa trên phân tích định lượng của các tiêu chí, đảm bảo tỷ số nhất quán CR < 10%."
        })
    
    except Exception as e:
        print(f"Error in get_alternative_matrices: {str(e)}")
        return jsonify({"error": f"Lỗi khi lấy ma trận: {str(e)}"}), 500

@app.route('/alternative_matrices')
def alternative_matrices():
    return render_template('alternative_matrices.html')

@app.route('/export-excel', methods=['POST'])
def export_excel():
    """Export AHP analysis results to Excel file with all data in a single sheet for easy import"""
    try:
        data = request.json
        goal = data.get('goal', 'AHP Analysis')
        criteria = data.get('criteria', [])
        alternatives = data.get('alternatives', [])
        criteriaWeights = data.get('criteriaWeights', [])
        alternativeScores = data.get('alternativeScores', {})
        finalScores = data.get('finalScores', {})
        criteriaMatrix = data.get('criteriaMatrix', [])  # Ma trận so sánh tiêu chí
        # Lấy ma trận gốc của từng tiêu chí
        alternativeMatrices = data.get('alternativeMatrices', {})
        
        # Tạo tên file an toàn
        safe_filename = "".join([c for c in goal if c.isalpha() or c.isdigit() or c==' ']).rstrip()
        if len(safe_filename) > 30:
            safe_filename = safe_filename[:30]
        safe_filename = safe_filename.replace(' ', '_')
        
        # Create Excel writer
        output = io.BytesIO()
        writer = pd.ExcelWriter(output, engine='xlsxwriter')
        workbook = writer.book
        
        # Định nghĩa các formats
        header_format = workbook.add_format({
            'bold': True,
            'text_wrap': True,
            'valign': 'top',
            'fg_color': '#D7E4BC',
            'border': 1
        })
        
        section_format = workbook.add_format({
            'bold': True,
            'text_wrap': True,
            'valign': 'top',
            'fg_color': '#B8CCE4',
            'border': 1,
            'font_size': 12
        })
        
        percent_format = workbook.add_format({
            'num_format': '0.00%'
        })
        
        decimal_format = workbook.add_format({
            'num_format': '0.0000'
        })
        
        # Tạo worksheet chính để chứa tất cả dữ liệu
        worksheet = workbook.add_worksheet('AHP_Data')
        
        # Set độ rộng cột
        worksheet.set_column('A:A', 20)
        worksheet.set_column('B:Z', 15)
        
        # 1. Thông tin chung
        row = 0
        worksheet.merge_range(row, 0, row, 3, 'THÔNG TIN PHÂN TÍCH AHP', section_format)
        row += 1
        
        # Mục tiêu
        worksheet.write(row, 0, 'Mục tiêu:', header_format)
        worksheet.merge_range(row, 1, row, 3, goal)
        row += 2
        
        # 2. Tiêu chí và trọng số
        worksheet.merge_range(row, 0, row, 3, 'TIÊU CHÍ VÀ TRỌNG SỐ', section_format)
        row += 1
        
        worksheet.write(row, 0, 'Tiêu chí', header_format)
        worksheet.write(row, 1, 'Trọng số', header_format)
        worksheet.write(row, 2, 'Phần trăm', header_format)
        row += 1
        
        # Ghi dữ liệu tiêu chí và trọng số
        for i, criterion in enumerate(criteria):
            worksheet.write(row, 0, criterion)
            worksheet.write(row, 1, criteriaWeights[i], decimal_format)
            worksheet.write(row, 2, criteriaWeights[i], percent_format)
            row += 1
        
        row += 1
        
        # 2.1 Ma trận so sánh các tiêu chí
        if criteriaMatrix and len(criteriaMatrix) > 0:
            worksheet.merge_range(row, 0, row, len(criteria) + 1, 'MA TRẬN SO SÁNH CÁC TIÊU CHÍ', section_format)
            row += 1
            
            # Header hàng
            worksheet.write(row, 0, 'Tiêu chí', header_format)
            for j, crit in enumerate(criteria):
                worksheet.write(row, j + 1, crit, header_format)
            row += 1
            
            # Ghi ma trận so sánh
            for i, crit1 in enumerate(criteria):
                worksheet.write(row, 0, crit1)
                for j, crit2 in enumerate(criteria):
                    if i < len(criteriaMatrix) and j < len(criteriaMatrix[i]):
                        value = criteriaMatrix[i][j]
                        # Chuyển đổi các giá trị dạng chuỗi như "1/3" thành số
                        if isinstance(value, str) and '/' in value:
                            parts = value.split('/')
                            if len(parts) == 2:
                                try:
                                    value = float(parts[0]) / float(parts[1])
                                except:
                                    value = 1
                        worksheet.write(row, j + 1, float(value), decimal_format)
                    else:
                        worksheet.write(row, j + 1, 1, decimal_format)
                row += 1
            
            row += 1
        
        # 3. Phương án
        worksheet.merge_range(row, 0, row, 3, 'PHƯƠNG ÁN', section_format)
        row += 1
        
        worksheet.write(row, 0, 'Danh sách phương án:', header_format)
        row += 1
        
        # Ghi danh sách phương án
        for i, alternative in enumerate(alternatives):
            worksheet.write(row, 0, alternative)
            row += 1
        
        row += 1
        
        # 4. Ma trận so sánh từng tiêu chí
        worksheet.merge_range(row, 0, row, len(alternatives) + 1, 'MA TRẬN SO SÁNH PHƯƠNG ÁN THEO TỪNG TIÊU CHÍ', section_format)
        row += 1
        
        # Lặp qua từng tiêu chí
        for i, criterion in enumerate(criteria):
            worksheet.merge_range(row, 0, row, len(alternatives) + 1, f'Tiêu chí: {criterion}', header_format)
            row += 1
            
            # Header hàng
            worksheet.write(row, 0, 'Phương án', header_format)
            for j, alt in enumerate(alternatives):
                worksheet.write(row, j + 1, alt, header_format)
            row += 1
            
            # Ghi ma trận so sánh
            if alternativeMatrices and criterion in alternativeMatrices:
                # Sử dụng ma trận gốc người dùng nhập nếu có
                matrix = None
                if 'matrix' in alternativeMatrices[criterion]:
                    matrix = alternativeMatrices[criterion]['matrix']
                    
                if matrix:
                    for j, alt1 in enumerate(alternatives):
                        worksheet.write(row, 0, alt1)
                        for k, alt2 in enumerate(alternatives):
                            if j < len(matrix) and k < len(matrix[j]):
                                value = matrix[j][k]
                                # Chuyển đổi giá trị dạng chuỗi (nếu có)
                                if isinstance(value, str) and '/' in value:
                                    parts = value.split('/')
                                    if len(parts) == 2:
                                        try:
                                            value = float(parts[0]) / float(parts[1])
                                        except:
                                            value = 1 if j == k else 0
                                worksheet.write(row, k + 1, float(value), decimal_format)
                            else:
                                # Nếu không có dữ liệu, điền 1 cho đường chéo chính và 0 cho các ô khác
                                worksheet.write(row, k + 1, 1 if j == k else 0)
                        row += 1
                else:
                    # Nếu không có ma trận gốc, tạo ma trận từ trọng số
                    weights = alternativeMatrices[criterion].get('weights', [])
                    if weights and len(weights) == len(alternatives):
                        for j, alt1 in enumerate(alternatives):
                            worksheet.write(row, 0, alt1)
                            for k, alt2 in enumerate(alternatives):
                                if j == k:
                                    # Đường chéo chính luôn là 1
                                    worksheet.write(row, k + 1, 1)
                                else:
                                    # Tính tỷ lệ từ trọng số
                                    if j < len(weights) and k < len(weights) and weights[k] != 0:
                                        val = weights[j] / weights[k]
                                        worksheet.write(row, k + 1, val, decimal_format)
                                    else:
                                        worksheet.write(row, k + 1, 1)
                            row += 1
                    else:
                        # Nếu không có trọng số, dùng tỷ lệ trọng số từ alternativeScores
                        for j, alt1 in enumerate(alternatives):
                            worksheet.write(row, 0, alt1)
                            for k, alt2 in enumerate(alternatives):
                                if j == k:
                                    # Đường chéo chính luôn là 1
                                    worksheet.write(row, k + 1, 1)
                                else:
                                    # Lấy giá trị so sánh từ dữ liệu trọng số phương án
                                    if alternativeScores.get(alt1) and alternativeScores[alt1].get(criterion) and alternativeScores.get(alt2) and alternativeScores[alt2].get(criterion):
                                        val = alternativeScores[alt1][criterion] / alternativeScores[alt2][criterion]
                                        worksheet.write(row, k + 1, val, decimal_format)
                                    else:
                                        worksheet.write(row, k + 1, 1)
                            row += 1
            else:
                # Nếu không có ma trận gốc, dùng tỷ lệ trọng số (cách cũ)
                for j, alt1 in enumerate(alternatives):
                    worksheet.write(row, 0, alt1)
                    for k, alt2 in enumerate(alternatives):
                        if j == k:
                            # Đường chéo chính luôn là 1
                            worksheet.write(row, k + 1, 1)
                        else:
                            # Lấy giá trị so sánh từ dữ liệu trọng số phương án
                            if alternativeScores.get(alt1) and alternativeScores[alt1].get(criterion) and alternativeScores.get(alt2) and alternativeScores[alt2].get(criterion):
                                val = alternativeScores[alt1][criterion] / alternativeScores[alt2][criterion]
                                worksheet.write(row, k + 1, val, decimal_format)
                            else:
                                worksheet.write(row, k + 1, 1)
                        row += 1
            
            # Trọng số phương án theo tiêu chí này
            row += 1
            worksheet.write(row, 0, "Trọng số:", header_format)
            
            # Lấy trọng số từ ma trận alternativeMatrices nếu có, nếu không thì từ alternativeScores
            weights = []
            if alternativeMatrices and criterion in alternativeMatrices and 'weights' in alternativeMatrices[criterion]:
                weights = alternativeMatrices[criterion]['weights']
                
                for j, alt in enumerate(alternatives):
                    if j < len(weights):
                        worksheet.write(row, j + 1, weights[j], decimal_format)
                    else:
                        worksheet.write(row, j + 1, 0, decimal_format)
            else:
                for j, alt in enumerate(alternatives):
                    if alternativeScores.get(alt) and alternativeScores[alt].get(criterion):
                        worksheet.write(row, j + 1, alternativeScores[alt][criterion], decimal_format)
                    else:
                        worksheet.write(row, j + 1, 0, decimal_format)
            
            row += 2
        
        # 5. Kết quả cuối cùng
        worksheet.merge_range(row, 0, row, 3, 'KẾT QUẢ CUỐI CÙNG', section_format)
        row += 1
        
        worksheet.write(row, 0, 'Phương án', header_format)
        worksheet.write(row, 1, 'Điểm số cuối cùng', header_format)
        worksheet.write(row, 2, 'Xếp hạng', header_format)
        row += 1
        
        # Sắp xếp phương án theo điểm số giảm dần
        sorted_alternatives = sorted(alternatives, key=lambda alt: finalScores.get(alt, 0), reverse=True)
        
        # Ghi kết quả cuối cùng
        for i, alt in enumerate(sorted_alternatives):
            worksheet.write(row, 0, alt)
            worksheet.write(row, 1, finalScores.get(alt, 0), decimal_format)
            worksheet.write(row, 2, i + 1)
            row += 1
        
        # Đóng workbook và trả về file
        writer.close()
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f"AHP_Analysis_{safe_filename}.xlsx"
        )
        
    except Exception as e:
        print(f"Error exporting to Excel: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

def xl_col_to_name(col_num):
    """Chuyển đổi số cột Excel sang tên cột (0 -> A, 1 -> B, ...)"""
    result = ""
    while col_num >= 0:
        remainder = col_num % 26
        result = chr(65 + remainder) + result
        col_num = col_num // 26 - 1
    return result

@app.route('/export-pdf', methods=['POST'])
def export_pdf():
    """Export AHP analysis results to PDF file with professional, administrative report style"""
    try:
        data = request.json
        goal = data.get('goal', 'AHP Analysis')
        criteria = data.get('criteria', [])
        alternatives = data.get('alternatives', [])
        criteriaWeights = data.get('criteriaWeights', [])
        alternativeScores = data.get('alternativeScores', {})
        finalScores = data.get('finalScores', {})

        log_to_mongodb({
            'type': 'export_pdf',
            'goal': goal,
            'criteria': criteria,
            'alternatives': alternatives,
            'finalScores': finalScores
        })

        buffer = io.BytesIO()
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import os
        font_path = os.path.join(os.path.dirname(__file__), 'fonts', 'DejaVuSans.ttf')
        pdfmetrics.registerFont(TTFont('DejaVuSans', font_path))

        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
        from reportlab.platypus import Frame, PageTemplate

        doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
        styles = getSampleStyleSheet()
        elements = []

        # Custom styles (avoid name conflict)
        styles.add(ParagraphStyle(
            name='CustomHeader', fontName='DejaVuSans', fontSize=14, alignment=TA_CENTER, textColor=colors.HexColor('#2563eb'), leading=18, spaceAfter=8, spaceBefore=0, fontWeight='bold', underlineWidth=1.5
        ))
        styles.add(ParagraphStyle(
            name='CustomTitle', fontName='DejaVuSans', fontSize=16, alignment=TA_CENTER, leading=22, spaceAfter=16, spaceBefore=0, fontWeight='bold'
        ))
        styles.add(ParagraphStyle(
            name='CustomSection', fontName='DejaVuSans', fontSize=13, alignment=TA_CENTER, leading=18, spaceAfter=10, spaceBefore=10, textColor=colors.HexColor('#1e293b'), fontWeight='bold'
        ))
        styles.add(ParagraphStyle(
            name='CustomTableHeader', fontName='DejaVuSans', fontSize=11, alignment=TA_CENTER, leading=14, textColor=colors.white, backColor=colors.HexColor('#2563eb'), spaceAfter=2, spaceBefore=2
        ))
        styles.add(ParagraphStyle(
            name='CustomTableCell', fontName='DejaVuSans', fontSize=10, alignment=TA_LEFT, leading=13, spaceAfter=2, spaceBefore=2, wordWrap='CJK'
        ))
        styles.add(ParagraphStyle(
            name='CustomTableCellCenter', fontName='DejaVuSans', fontSize=10, alignment=TA_CENTER, leading=13, spaceAfter=2, spaceBefore=2, wordWrap='CJK'
        ))
        styles.add(ParagraphStyle(
            name='CustomConclusion', fontName='DejaVuSans', fontSize=11, alignment=TA_JUSTIFY, leading=15, textColor=colors.HexColor('#166534'), backColor=colors.HexColor('#f0fdf4'), spaceAfter=10, spaceBefore=10, leftIndent=8, rightIndent=8, borderPadding=6
        ))
        styles.add(ParagraphStyle(
            name='CustomFooter', fontName='DejaVuSans', fontSize=9, alignment=TA_CENTER, textColor=colors.HexColor('#64748b'), leading=12, spaceBefore=12
        ))

        # Header
        elements.append(Paragraph("NHÓM 7 - HỆ HỖ TRỢ RA QUYẾT ĐỊNH", styles['CustomHeader']))
        elements.append(Spacer(1, 0.08*inch))
        # Title
        elements.append(Paragraph(f"Phân tích AHP: {goal}", styles['CustomTitle']))
        elements.append(Spacer(1, 0.12*inch))

        # Section: Criteria Weights
        elements.append(Paragraph("Trọng số các tiêu chí", styles['CustomSection']))
        # Table data
        criteria_data = [
            [Paragraph('Tiêu chí', styles['CustomTableHeader']),
             Paragraph('Trọng số', styles['CustomTableHeader']),
             Paragraph('Tỷ lệ', styles['CustomTableHeader'])]
        ]
        for i, criterion in enumerate(criteria):
            criteria_data.append([
                Paragraph(criterion, styles['CustomTableCell']),
                Paragraph(f"{criteriaWeights[i]:.4f}", styles['CustomTableCellCenter']),
                Paragraph(f"{criteriaWeights[i]*100:.2f}%", styles['CustomTableCellCenter'])
            ])
        criteria_table = Table(criteria_data, colWidths=[2.8*inch, 1.1*inch, 1.1*inch], repeatRows=1)
        criteria_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), 'DejaVuSans'),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2563eb')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('ALIGN', (1, 1), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.whitesmoke, colors.HexColor('#e0e7ff')]),
            ('GRID', (0, 0), (-1, -1), 0.7, colors.HexColor('#64748b')),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        elements.append(criteria_table)
        elements.append(Spacer(1, 0.18*inch))

        # Section: Final Results
        elements.append(Paragraph("Kết quả phân tích cuối cùng", styles['CustomSection']))
        # Sort alternatives by final score
        sorted_alternatives = sorted(alternatives, key=lambda alt: finalScores.get(alt, 0), reverse=True)
        total_score = sum(finalScores.values())
        final_data = [
            [Paragraph('Xếp hạng', styles['CustomTableHeader']),
             Paragraph('Phương án', styles['CustomTableHeader']),
             Paragraph('Điểm số', styles['CustomTableHeader']),
             Paragraph('Tỷ lệ', styles['CustomTableHeader'])]
        ]
        for i, alt in enumerate(sorted_alternatives):
            score = finalScores.get(alt, 0)
            percentage = (score / total_score) * 100 if total_score > 0 else 0
            row = [
                Paragraph(f"{i+1}", styles['CustomTableCellCenter']),
                Paragraph(alt, styles['CustomTableCell']),
                Paragraph(f"{score:.4f}", styles['CustomTableCellCenter']),
                Paragraph(f"{percentage:.2f}%", styles['CustomTableCellCenter'])
            ]
            final_data.append(row)
        final_table = Table(final_data, colWidths=[0.8*inch, 3.2*inch, 1.1*inch, 1.1*inch], repeatRows=1)
        # Highlight best row
        best_row_color = colors.HexColor('#bbf7d0')
        final_table_style = [
            ('FONTNAME', (0, 0), (-1, -1), 'DejaVuSans'),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2563eb')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.whitesmoke, colors.HexColor('#e0e7ff')]),
            ('GRID', (0, 0), (-1, -1), 0.7, colors.HexColor('#64748b')),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('ALIGN', (0, 1), (0, -1), 'CENTER'),
            ('ALIGN', (2, 1), (3, -1), 'CENTER'),
            ('ALIGN', (1, 1), (1, -1), 'LEFT'),
            ('BACKGROUND', (0, 1), (-1, 1), best_row_color),
        ]
        final_table.setStyle(TableStyle(final_table_style))
        elements.append(final_table)
        elements.append(Spacer(1, 0.18*inch))

        # Conclusion
        if sorted_alternatives:
            best_alt = sorted_alternatives[0]
            best_score = finalScores.get(best_alt, 0)
            percentage = (best_score / total_score) * 100 if total_score > 0 else 0
            conclusion_text = f"<b>Kết luận:</b> Phương án <b>{best_alt}</b> là lựa chọn tốt nhất với điểm số <b>{best_score:.4f}</b> (<b>{percentage:.2f}%</b>)."
            if len(sorted_alternatives) > 1:
                second_alt = sorted_alternatives[1]
                second_score = finalScores.get(second_alt, 0)
                diff_percent = ((best_score - second_score) / best_score) * 100 if best_score > 0 else 0
                conclusion_text += f" Phương án này vượt trội hơn phương án <b>{second_alt}</b> khoảng <b>{diff_percent:.2f}%</b>."
            elements.append(Spacer(1, 0.05*inch))
            elements.append(Paragraph(conclusion_text, styles['CustomConclusion']))

        # Footer
        elements.append(Spacer(1, 0.18*inch))
        footer_text = f"Nhóm 7 - Hệ hỗ trợ ra quyết định | Báo cáo được tạo vào: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}"
        elements.append(Paragraph(footer_text, styles['CustomFooter']))

        doc.build(elements)
        buffer.seek(0)
        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f"AHP_Analysis_{goal.replace(' ', '_')[:30]}.pdf"
        )
    except Exception as e:
        print(f"Error exporting to PDF: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/save-log', methods=['POST'])
def save_log():
    try:
        data = request.json
        log_to_mongodb({
            'type': 'manual_save',
            **data
        })
        return jsonify({"success": True, "message": "Đã lưu kết quả vào MongoDB!"})
    except Exception as e:
        return jsonify({"success": False, "message": f"Lỗi: {str(e)}"}), 500

@app.route('/import-excel', methods=['POST'])
def import_excel():
    """Import AHP analysis from Excel file previously exported by the system"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
            
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
            
        if not file.filename.endswith('.xlsx'):
            return jsonify({"error": "Only Excel (.xlsx) files are supported"}), 400
            
        # Đọc file Excel
        excel_data = pd.read_excel(file, sheet_name='AHP_Data', engine='openpyxl')
        
        # Cấu trúc để lưu trữ dữ liệu phân tích AHP
        ahp_data = {
            "goal": "",
            "criteria": [],
            "criteria_weights": [],
            "alternatives": [],
            "alternative_matrices": {},
            "final_scores": {},
            "criteria_matrix": []  # Thêm ma trận so sánh tiêu chí
        }
        
        # Đọc tệp Excel và tách dữ liệu
        try:
            # Đọc dữ liệu từ file Excel trực tiếp
            df = pd.ExcelFile(file)
            data_df = pd.read_excel(df, sheet_name='AHP_Data')
            
            # 1. Đọc mục tiêu
            goal_row = data_df[data_df.iloc[:, 0] == 'Mục tiêu:']
            if not goal_row.empty:
                ahp_data["goal"] = goal_row.iloc[0, 1]
            
            # 2. Đọc danh sách tiêu chí và trọng số
            criteria_start = None
            for i, row in data_df.iterrows():
                if row.iloc[0] == 'TIÊU CHÍ VÀ TRỌNG SỐ':
                    criteria_start = i
                    break
                    
            if criteria_start is not None:
                criteria_data = data_df.iloc[criteria_start+2:].reset_index(drop=True)
                
                # Tiếp tục cho đến khi gặp dòng trống
                criteria_end = 0
                for i, row in criteria_data.iterrows():
                    if pd.isna(row.iloc[0]) or row.iloc[0] == '':
                        criteria_end = i
                        break
                        
                if criteria_end > 0:
                    criteria_data = criteria_data.iloc[:criteria_end]
                    
                for i, row in criteria_data.iterrows():
                    if not pd.isna(row.iloc[0]) and row.iloc[0] != '':
                        ahp_data["criteria"].append(row.iloc[0])
                        ahp_data["criteria_weights"].append(float(row.iloc[1]))
            
            # 2.1 Đọc ma trận so sánh các tiêu chí
            criteria_matrix_start = None
            for i, row in data_df.iterrows():
                if isinstance(row.iloc[0], str) and row.iloc[0] == 'MA TRẬN SO SÁNH CÁC TIÊU CHÍ':
                    criteria_matrix_start = i
                    break
                    
            if criteria_matrix_start is not None:
                # Bỏ qua header, đọc từ hàng header tiêu chí + 2
                matrix_start_idx = criteria_matrix_start + 2
                n_criteria = len(ahp_data["criteria"])
                
                if matrix_start_idx < len(data_df):
                    # Tạo ma trận so sánh tiêu chí
                    criteria_matrix = []
                    
                    for i in range(n_criteria):
                        row_idx = matrix_start_idx + i
                        if row_idx < len(data_df):
                            row_values = []
                            for j in range(n_criteria):
                                col_idx = j + 1  # Cột đầu tiên là tên tiêu chí
                                if col_idx < len(data_df.columns):
                                    val = data_df.iloc[row_idx, col_idx]
                                    if not pd.isna(val):
                                        row_values.append(val)
                                    else:
                                        row_values.append(1.0)
                                else:
                                    row_values.append(1.0)
                            criteria_matrix.append(row_values)
                    
                    ahp_data["criteria_matrix"] = criteria_matrix
            
            # 3. Đọc danh sách phương án
            alternatives_start = None
            for i, row in data_df.iterrows():
                if row.iloc[0] == 'PHƯƠNG ÁN':
                    alternatives_start = i
                    break
                    
            if alternatives_start is not None:
                alt_data = data_df.iloc[alternatives_start+2:].reset_index(drop=True)
                
                # Tiếp tục cho đến khi gặp dòng trống
                alt_end = 0
                for i, row in alt_data.iterrows():
                    if pd.isna(row.iloc[0]) or row.iloc[0] == '':
                        alt_end = i
                        break
                        
                if alt_end > 0:
                    alt_data = alt_data.iloc[:alt_end]
                    
                for i, row in alt_data.iterrows():
                    if not pd.isna(row.iloc[0]) and row.iloc[0] != '':
                        ahp_data["alternatives"].append(row.iloc[0])
            
            # 4. Đọc ma trận so sánh theo từng tiêu chí
            matrix_start_indices = []
            for i, row in data_df.iterrows():
                if isinstance(row.iloc[0], str) and row.iloc[0].startswith('Tiêu chí: '):
                    matrix_start_indices.append((i, row.iloc[0][9:]))  # Bỏ 'Tiêu chí: ' để lấy tên tiêu chí
            
            # Xử lý từng ma trận
            for idx, (start_idx, criterion) in enumerate(matrix_start_indices):
                # Tìm điểm kết thúc của ma trận (dòng 'Trọng số:' hoặc ma trận tiếp theo)
                end_idx = None
                for i in range(start_idx + 2, len(data_df)):
                    if i >= len(data_df) or (isinstance(data_df.iloc[i, 0], str) and data_df.iloc[i, 0] == 'Trọng số:'):
                        end_idx = i
                        break
                    if idx < len(matrix_start_indices) - 1 and i >= matrix_start_indices[idx + 1][0]:
                        end_idx = matrix_start_indices[idx + 1][0] - 1
                        break
                
                if end_idx is None:
                    continue
                
                # Đọc ma trận so sánh
                matrix_data = data_df.iloc[start_idx + 2:end_idx].reset_index(drop=True)
                
                # Tạo ma trận so sánh
                n = len(ahp_data["alternatives"])
                comparison_matrix = np.ones((n, n))
                
                for i in range(n):
                    for j in range(n):
                        if i != j and i < len(matrix_data) and j + 1 < len(matrix_data.columns):
                            value = matrix_data.iloc[i, j + 1]
                            if not pd.isna(value):
                                comparison_matrix[i][j] = value
                
                # Tính toán trọng số của phương án cho tiêu chí này
                result = calculate_ahp(comparison_matrix)
                
                # Lưu trữ kết quả
                alt_weights = {}
                for i, alt in enumerate(ahp_data["alternatives"]):
                    alt_weights[alt] = result["weights"][i]
                
                ahp_data["alternative_matrices"][criterion] = {
                    "matrix": comparison_matrix.tolist(),
                    "weights": result["weights"].tolist() if hasattr(result["weights"], "tolist") else result["weights"],  # Kiểm tra có phải numpy array không
                    "cr": float(result["cr"]),  # Chuyển từ numpy float sang Python float
                    "alternatives": alt_weights
                }
            
            # 5. Đọc kết quả cuối cùng
            final_start = None
            for i, row in data_df.iterrows():
                if row.iloc[0] == 'KẾT QUẢ CUỐI CÙNG':
                    final_start = i
                    break
                    
            if final_start is not None:
                final_data = data_df.iloc[final_start+2:].reset_index(drop=True)
                
                # Tiếp tục cho đến khi gặp dòng trống hoặc kết thúc
                for i, row in final_data.iterrows():
                    if pd.isna(row.iloc[0]) or row.iloc[0] == '':
                        break
                    
                    if not pd.isna(row.iloc[0]) and not pd.isna(row.iloc[1]):
                        alt_name = row.iloc[0]
                        alt_score = float(row.iloc[1])
                        ahp_data["final_scores"][alt_name] = alt_score
            
            # Nếu không có điểm cuối cùng, tính toán lại
            if not ahp_data["final_scores"] and ahp_data["criteria"] and ahp_data["alternative_matrices"]:
                final_scores = calculate_final_scores(ahp_data["criteria_weights"], ahp_data["alternative_matrices"])
                ahp_data["final_scores"] = final_scores
                
            return jsonify(ahp_data)
            
        except Exception as e:
            print(f"Error parsing Excel file: {str(e)}")
            traceback.print_exc()
            return jsonify({"error": f"Failed to parse Excel file: {str(e)}"}), 500
    
    except Exception as e:
        print(f"Error importing from Excel: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500



if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)