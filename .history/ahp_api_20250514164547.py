from fastapi import FastAPI, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional, Dict
import uvicorn
import os
import json
import requests
import logging
import re
import tempfile
from tenacity import retry, stop_after_attempt, wait_exponential

# Thêm thư viện xử lý file
import io
import docx
import PyPDF2  # Thư viện xử lý PDF

# Cấu hình logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Constants
API_URL = os.getenv('API_URL', "http://localhost:1234/v1/chat/completions")
MAX_RETRIES = 3
TIMEOUT = 30
MAX_FILE_CONTENT_LENGTH = 100000
DEFAULT_MAX_CRITERIA = 4
DEFAULT_MAX_ALTERNATIVES = 4

app = FastAPI(
    title="AHP Suggestion API",
    description="API for getting AHP criteria and alternatives suggestions",
    version="1.0.0"
)

# Cấu hình CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class AHPSuggestionResponse(BaseModel):
    criteria: List[str]
    alternatives: List[str]

def read_docx_file(file_content: bytes) -> str:
    """
    Đọc nội dung file DOCX và chuyển thành string.
    
    Args:
        file_content: Nội dung file dạng bytes
    
    Returns:
        str: Nội dung text từ file DOCX
    """
    try:
        # Lưu file tạm để xử lý
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(file_content)
            temp_path = temp_file.name
        
        # Đọc nội dung từ file docx
        doc = docx.Document(temp_path)
        content = []
        
        # Trích xuất text từ đoạn văn
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        # Trích xuất text từ bảng
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                if any(cell for cell in row_text):  # Chỉ thêm hàng có nội dung
                    content.append(' | '.join(row_text))
        
        # Xóa file tạm
        os.unlink(temp_path)
        
        return '\n'.join(content)
    except Exception as e:
        logger.error(f"Error reading DOCX: {str(e)}")
        raise ValueError(f"Lỗi khi đọc file DOCX: {str(e)}")

def read_pdf_file(file_content: bytes) -> str:
    """
    Đọc nội dung file PDF và chuyển thành string.
    
    Args:
        file_content: Nội dung file dạng bytes
    
    Returns:
        str: Nội dung text từ file PDF
    """
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = []
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        return '\n'.join(text)
    except Exception as e:
        logger.error(f"Error reading PDF: {str(e)}")
        raise ValueError(f"Lỗi khi đọc file PDF: {str(e)}")

async def extract_text_from_file(file: UploadFile) -> str:
    """
    Trích xuất text từ các loại file khác nhau.
    
    Args:
        file: File được upload
        
    Returns:
        str: Nội dung text từ file
    """
    file_content = await file.read()
    file_extension = file.filename.lower().split('.')[-1]
    
    try:
        if file_extension == 'txt':
            # Xử lý text file
            return file_content.decode('utf-8')
            
        elif file_extension == 'docx':
            # Xử lý Word file
            return read_docx_file(file_content)
            
        elif file_extension == 'pdf':
            # Xử lý PDF file
            return read_pdf_file(file_content)
            
        else:
            raise ValueError(f"Không hỗ trợ định dạng file: {file_extension}. Chỉ hỗ trợ .txt, .docx và .pdf")
            
    except Exception as e:
        logger.error(f"Error extracting text from file: {str(e)}")
        raise ValueError(f"Không thể đọc nội dung file: {str(e)}")

def validate_goal(goal: str) -> bool:
    """Kiểm tra tính hợp lệ của goal."""
    return bool(goal and goal.strip())

def validate_response(result: dict) -> bool:
    """Kiểm tra tính hợp lệ của response."""
    return (
        isinstance(result, dict) and
        'criteria' in result and
        'alternatives' in result and
        all(isinstance(c, str) for c in result['criteria']) and
        all(isinstance(a, str) for a in result['alternatives'])
    )

def create_prompt(goal: str, file_content: Optional[str] = None) -> str:
    """
    Tạo prompt cho LLM dựa trên input.
    
    Args:
        goal: Mục tiêu cần phân tích
        file_content: Nội dung file (optional)
    
    Returns:
        str: Prompt được tạo
    """
    # Sử dụng prompt từ app.py vì đã được tối ưu
    if file_content:
        prompt = f"""Hãy phân tích mục tiêu và nội dung file đính kèm, sau đó đề xuất {DEFAULT_MAX_CRITERIA} tiêu chí đánh giá quan trọng nhất và {DEFAULT_MAX_ALTERNATIVES} phương án lựa chọn phù hợp nhất với mục tiêu đã cho:
        
        Mục tiêu: {goal}
        
        Lưu ý:
        - Ưu tiên các tiêu chí mang tính chất định lượng 
        - Chỉ đề xuất {DEFAULT_MAX_CRITERIA} tiêu chí quan trọng nhất
        - Chỉ đề xuất {DEFAULT_MAX_ALTERNATIVES} phương án lựa chọn tốt nhất
        - Đảm bảo các tiêu chí và phương án phải thực tế và có thể so sánh được
        - Phân tích cả nội dung file để đề xuất cho phù hợp
        
        QUAN TRỌNG: Chỉ trả về kết quả theo định dạng JSON như sau, không thêm text khác:
        {{
            "criteria": ["tiêu chí 1", "tiêu chí 2", "tiêu chí 3", "tiêu chí 4"],
            "alternatives": ["phương án 1", "phương án 2", "phương án 3", "phương án 4"]
        }}
        
        Dữ liệu như sau:
        {file_content[:MAX_FILE_CONTENT_LENGTH]}
        """
    else:
        prompt = f"""Hãy phân tích mục tiêu sau và đề xuất {DEFAULT_MAX_CRITERIA} tiêu chí đánh giá quan trọng nhất và {DEFAULT_MAX_ALTERNATIVES} phương án lựa chọn phù hợp nhất:
        
        Mục tiêu: {goal}
        
        Lưu ý:
        - Chỉ đề xuất {DEFAULT_MAX_CRITERIA} tiêu chí quan trọng nhất
        - Chỉ đề xuất {DEFAULT_MAX_ALTERNATIVES} phương án lựa chọn tốt nhất
        - Đảm bảo các tiêu chí và phương án phải thực tế và có thể so sánh được
        - Ưu tiên các tiêu chí mang tính chất định lượng
        
        QUAN TRỌNG: Chỉ trả về kết quả theo định dạng JSON như sau, không thêm text khác:
        {{
            "criteria": ["tiêu chí 1", "tiêu chí 2", "tiêu chí 3", "tiêu chí 4"],
            "alternatives": ["phương án 1", "phương án 2", "phương án 3", "phương án 4"]
        }}
        """
    
    return prompt

@retry(
    stop=stop_after_attempt(MAX_RETRIES),
    wait=wait_exponential(multiplier=1, min=4, max=10)
)
async def call_llm_api(prompt: str, has_file: bool) -> dict:
    """
    Gọi LLM API với cơ chế retry.
    
    Args:
        prompt: Prompt để gửi đến LLM
        has_file: Có file đính kèm hay không
    
    Returns:
        dict: Response từ LLM API
        
    Raises:
        HTTPException: Khi có lỗi từ API
    """
    headers = {
        "Content-Type": "application/json"
    }
    
    data = {
        "messages": [
            {
                "role": "system",
                "content": "Bạn là chuyên gia phân tích và tư vấn ra quyết định. Hãy trả về kết quả chính xác theo format JSON được yêu cầu."
            },
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.7,
        "max_tokens": 800 if has_file else 500
    }

    try:
        response = requests.post(
            API_URL,
            headers=headers,
            json=data,
            timeout=TIMEOUT
        )
        response.raise_for_status()
        return response.json()
    except requests.exceptions.RequestException as e:
        logger.error(f"API call failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Lỗi khi gọi LLM API: {str(e)}")

def process_llm_response(response: dict) -> dict:
    """
    Xử lý và validate response từ LLM.
    
    Args:
        response: Response từ LLM API
    
    Returns:
        dict: Kết quả đã được xử lý
        
    Raises:
        ValueError: Khi response không hợp lệ
    """
    if 'choices' not in response or not response['choices']:
        raise ValueError("Invalid response format from LLM")
        
    content = response['choices'][0]['message']['content']
    
    # Extract và parse JSON
    try:
        # Tìm đoạn JSON trong response
        start_idx = content.find('{')
        end_idx = content.rfind('}') + 1
        if start_idx != -1 and end_idx != -1:
            json_str = content[start_idx:end_idx]
            result = json.loads(json_str)
            
            # Validate và clean kết quả
            if not validate_response(result):
                raise ValueError("Invalid JSON structure")
            
            return {
                'criteria': [str(c).strip() for c in result['criteria'][:DEFAULT_MAX_CRITERIA]],
                'alternatives': [str(a).strip() for a in result['alternatives'][:DEFAULT_MAX_ALTERNATIVES]]
            }
        else:
            raise ValueError("No JSON found in response")
    except json.JSONDecodeError as e:
        logger.error(f"JSON parsing error: {str(e)}")
        raise ValueError(f"Cannot parse JSON from response: {str(e)}")

@app.post("/api/get-suggestions", response_model=AHPSuggestionResponse)
async def get_suggestions(
    goal: str = Form(...),
    file: Optional[UploadFile] = None
):
    """
    Endpoint để lấy gợi ý về tiêu chí và phương án từ LLM.
    
    Args:
        goal: Mục tiêu cần phân tích
        file: File đính kèm (optional)
    
    Returns:
        AHPSuggestionResponse: Kết quả gợi ý
        
    Raises:
        HTTPException: Khi có lỗi xảy ra
    """
    try:
        if not validate_goal(goal):
            raise ValueError("Mục tiêu không được để trống")

        # Đọc nội dung file nếu có
        file_content = None
        if file:
            try:
                file_content = await extract_text_from_file(file)
                logger.info(f"Đọc thành công file {file.filename} ({len(file_content)} bytes)")
            except ValueError as e:
                raise HTTPException(status_code=400, detail=str(e))

        # Tạo prompt và gọi API
        prompt = create_prompt(goal, file_content)
        response = await call_llm_api(prompt, file is not None)
        
        # Xử lý kết quả
        result = process_llm_response(response)
        
        return AHPSuggestionResponse(**result)

    except ValueError as e:
        logger.error(f"Validation error: {str(e)}")
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Unexpected error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000) 