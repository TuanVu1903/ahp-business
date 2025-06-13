from flask import Flask, render_template, request, jsonify, send_file
import numpy as np
import requests
import os
import io
import tempfile
import json
import traceback
# Thêm thư viện xử lý file docx
import docx
# Thêm thư viện xử lý PDF và Excel
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
# Import MongoDB
from pymongo import MongoClient
from datetime import datetime

app = Flask(__name__)

# Kết nối MongoDB
try:
    mongo_client = MongoClient('mongodb://localhost:27017/')
    db = mongo_client['ahp']
    logs_collection = db['logs']
    print("MongoDB connected successfully!")
except Exception as e:
    print(f"MongoDB connection error: {str(e)}")
    logs_collection = None

# Hàm ghi log vào MongoDB
def log_to_mongodb(data):
    try:
        if logs_collection is not None:
            # Thêm timestamp
            data['timestamp'] = datetime.now()
            # Thêm vào collection
            logs_collection.insert_one(data)
            return True
    except Exception as e:
        print(f"MongoDB log error: {str(e)}")
    return False

# Cấu hình API cho LLM
API_URL = os.getenv('API_URL', "http://localhost:1234/v1/chat/completions")


def read_docx_file(file_storage):
    """Đọc nội dung file DOCX và chuyển thành string."""
    try:
        if 'docx' not in globals():
            return "Thư viện python-docx chưa được cài đặt. Không thể đọc file DOCX."
        
        # Lưu file tạm để xử lý
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            file_storage.save(temp_file.name)
            temp_path = temp_file.name
        
        # Đọc nội dung từ file docx
        doc = docx.Document(temp_path)
        content = []
        
        # Trích xuất text từ đoạn văn
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        # Trích xuất text từ bảng
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                if any(cell for cell in row_text):  # Chỉ thêm hàng có nội dung
                    content.append(' | '.join(row_text))
        
        # Xóa file tạm
        os.unlink(temp_path)
        
        return '\n'.join(content)
    except Exception as e:
        print(f"Error reading DOCX: {str(e)}")
        return f"Lỗi khi đọc file DOCX: {str(e)}"

def get_llm_suggestions(goal, file_content=None):
    """
    Nhận goal và nội dung file để phân tích và đề xuất tiêu chí và phương án.
    
    Args:
        goal: Mục tiêu cần phân tích
        file_content: Nội dung file đã được đọc (dạng string)
    
    Returns:
        Dict chứa các tiêu chí và phương án được đề xuất
    """
    # Xác định nội dung prompt dựa trên việc có file hay không
    if file_content:
        prompt = f"""Hãy phân tích mục tiêu và nội dung file đính kèm, sau đó đề xuất 4 tiêu chí đánh giá quan trọng nhất và 4 phương án lựa chọn phù hợp nhất với mục tiêu đã cho:
        
        Mục tiêu: {goal}
        
        Lưu ý:
        - Ưu tiên các tiêu chí mang tính chất định lượng 
        - Chỉ đề xuất 4 tiêu chí quan trọng nhất
        - Chỉ đề xuất 4 phương án lựa chọn tốt nhất
        - Đảm bảo các tiêu chí và phương án phải thực tế và có thể so sánh được
        - Phân tích cả nội dung file để đề xuất cho phù hợp
        
        Hãy trả về kết quả theo định dạng JSON như sau:
        {{
            "criteria": ["tiêu chí 1", "tiêu chí 2", "tiêu chí 3", "tiêu chí 4"],
            "alternatives": ["phương án 1", "phương án 2", "phương án 3", "phương án 4"]
        }}
        Dữ liệu như sau:
        {file_content[:100000]}  
        tiêu chí phải phù hợp với mục tiêu {goal};
        """
    else:
        prompt = f"""Hãy phân tích mục tiêu sau và đề xuất 4 tiêu chí đánh giá quan trọng nhất và 5 phương án lựa chọn phù hợp nhất:
        Mục tiêu: {goal}
        
        Lưu ý:
        - Chỉ đề xuất 4 tiêu chí quan trọng nhất
        - Chỉ đề xuất 4 phương án lựa chọn tốt nhất
        - Đảm bảo các tiêu chí và phương án phải thực tế và có thể so sánh được
        
        Hãy trả về kết quả theo định dạng JSON như sau:
        {{
            "criteria": ["tiêu chí 1", "tiêu chí 2", "tiêu chí 3", "tiêu chí 4"],
            "alternatives": ["phương án 1", "phương án 2", "phương án 3", "phương án 4"]
        }}
        """
    try:
        response = requests.post(
            API_URL,
            json={
                "messages": [
                    {"role": "system", "content": "Bạn là một chuyên gia phân tích và tư vấn ra quyết định. Hãy phân tích mục tiêu và đề xuất 4 tiêu chí đánh giá quan trọng nhất và 5 phương án thay thế phù hợp nhất."},
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.7,
                "max_tokens": 800 if file_content else 500  # Tăng max_tokens nếu có file
            }
        )
        response.raise_for_status()
        
        response_data = response.json()
        if 'choices' in response_data and len(response_data['choices']) > 0:
            content = response_data['choices'][0]['message']['content']
            try:
                import json
                start_idx = content.find('{')
                end_idx = content.rfind('}') + 1
                if start_idx != -1 and end_idx != 0:
                    json_str = content[start_idx:end_idx]
                    result = json.loads(json_str)
                    result['criteria'] = result['criteria'][:5]
                    result['alternatives'] = result['alternatives'][:5]
                    return result
            except json.JSONDecodeError as e:
                print("Error parsing JSON from response:", e)
                return {"error": "Không thể xử lý kết quả từ LLM. Vui lòng thử lại."}
        
        return {"error": "Không nhận được kết quả phù hợp từ LLM."}
    except Exception as e:
        print(f"Error: {str(e)}")
        return {"error": f"Lỗi khi kết nối đến LLM: {str(e)}"}

def calculate_ahp(comparison_matrix):
    matrix_original = comparison_matrix.copy()
    n = comparison_matrix.shape[0]

    # Tính sum từng cột
    arr_sum = []
    for i in range(n):
        sum = 0
        for j in range(n):
            sum += comparison_matrix[j][i]
        arr_sum.append(sum)

    # Chuẩn hóa ma trận so sánh cặp
    normalized_matrix = np.zeros((n, n))
    for i in range(n):
        for j in range(n):
            normalized_matrix[i][j] = comparison_matrix[i][j] / arr_sum[j]

    # Tính Criteria Weight
    arr_avg = []
    for i in range(n):
        sum = 0
        for j in range(n):
            sum += normalized_matrix[i][j]
        arr_avg.append(sum/n)

    # Nhân Criteria Weight với ma trận so sánh cặp
    for i in range(n):
        for j in range(n):
            matrix_original[j][i] = matrix_original[j][i] * arr_avg[i]
    
    # Tính Weighted Sum Value
    arr_WSV = []
    for i in range(n):
        sum = 0
        for j in range(n):
            sum += matrix_original[i][j]
        arr_WSV.append(sum)
    
    # Tính Consistency Vector
    arr_CV = []
    for i in range(len(arr_WSV)):
        arr_CV.append(arr_WSV[i] / arr_avg[i])

    # Tính Lambda Max
    sum = 0
    for value in arr_CV:
        sum += value
    lambda_max = sum / len(arr_CV)

    # Tính CI
    ci = (lambda_max - n) / (n - 1)

    # RI values
    ri_values = [0.00, 0.00, 0.58, 0.90, 1.12, 1.24, 1.32, 1.41, 1.45, 1.49, 1.51, 1.54, 1.56, 1.58, 1.59]
    ri = ri_values[n-1] if n <= len(ri_values) else ri_values[-1]

    # Tính CR
    cr = ci / ri

    return {
        "weights": arr_avg,
        "lambda_max": float(lambda_max),
        "ci": float(ci),
        "cr": float(cr),
        "normalized_matrix": normalized_matrix.tolist()
    }

def calculate_final_scores(criteria_weights, alt_matrices_by_criteria):
    """
    Tính điểm cuối cùng cho các phương án dựa trên trọng số tiêu chí và ma trận đánh giá theo từng tiêu chí
    
    Args:
        criteria_weights: Danh sách trọng số các tiêu chí
        alt_matrices_by_criteria: Dict với key là tên tiêu chí và value là ma trận trọng số của phương án theo tiêu chí đó
    
    Returns:
        Dict chứa điểm số cuối cùng của từng phương án
    """
    final_scores = {}
    
    try:
        # Lấy danh sách các phương án từ ma trận đầu tiên
        first_criterion = list(alt_matrices_by_criteria.keys())[0]
        
        if 'alternatives' in alt_matrices_by_criteria[first_criterion]:
            # Dữ liệu từ frontend có cấu trúc là dict
            alternatives = list(alt_matrices_by_criteria[first_criterion]["alternatives"].keys())
        else:
            return {"error": "Định dạng ma trận không hợp lệ"}
        
        # Khởi tạo điểm số cho từng phương án
        for alt in alternatives:
            final_scores[alt] = 0
        
        # Tính tổng điểm cho từng phương án
        for i, criterion in enumerate(alt_matrices_by_criteria.keys()):
            # Kiểm tra chỉ số i có hợp lệ không
            if i >= len(criteria_weights):
                raise ValueError(f"Chỉ số tiêu chí không khớp. Có thể thiếu trọng số cho tiêu chí '{criterion}'")
            
            criterion_weight = criteria_weights[i]
            
            # Kiểm tra giá trị có phải số không
            if not isinstance(criterion_weight, (int, float)):
                raise ValueError(f"Trọng số tiêu chí '{criterion}' không phải là số: '{criterion_weight}'")
            
            # Kiểm tra cấu trúc dữ liệu
            if 'alternatives' not in alt_matrices_by_criteria[criterion]:
                raise ValueError(f"Không tìm thấy trọng số phương án cho tiêu chí '{criterion}'")
            
            alt_weights = alt_matrices_by_criteria[criterion]["alternatives"]
            
            for alt in alternatives:
                if alt not in alt_weights:
                    raise ValueError(f"Không tìm thấy trọng số cho phương án '{alt}' trong tiêu chí '{criterion}'")
                
                # Kiểm tra giá trị trọng số của phương án
                alt_weight = alt_weights[alt]
                if not isinstance(alt_weight, (int, float)):
                    raise ValueError(f"Trọng số của phương án '{alt}' trong tiêu chí '{criterion}' không phải là số: '{alt_weight}'")
                
                final_scores[alt] += criterion_weight * alt_weight
                
        return final_scores
        
    except Exception as e:
        print(f"Lỗi khi tính điểm cuối cùng: {str(e)}")
        raise ValueError(f"Không thể tính điểm cuối cùng: {str(e)}")

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/get_suggestions', methods=['POST'])
def get_suggestions():
    """Xử lý yêu cầu phân tích từ form hoặc AJAX."""
    try:
        # Lấy goal từ form hoặc JSON
        if request.form:
            goal = request.form.get('goal', '')
        else:
            goal = request.json.get('goal', '')
        
        # Kiểm tra goal có tồn tại
        if not goal:
            return jsonify({"error": "Vui lòng nhập mục tiêu"}), 400
        
        # Xử lý file nếu có
        file_content = None
        if 'file' in request.files and request.files['file'].filename != '':
            file = request.files['file']
            filename = file.filename.lower()
            
            # Kiểm tra loại file và đọc nội dung
            if filename.endswith('.docx'):
                file_content = read_docx_file(file)
            elif filename.endswith('.doc'):
                return jsonify({"error": "File .doc cũ không được hỗ trợ. Vui lòng chuyển đổi sang .docx"}), 400
            else:
                return jsonify({"error": "Chỉ hỗ trợ file Word (.docx)"}), 400
        
        # Gọi hàm get_llm_suggestions
        suggestions = get_llm_suggestions(goal, file_content)
        return jsonify(suggestions)
        
    except Exception as e:
        print(f"Error in get_suggestions: {str(e)}")
        return jsonify({"error": f"Lỗi xử lý yêu cầu: {str(e)}"}), 500

@app.route('/calculate_ahp', methods=['POST'])
def calculate():
    data = request.json
    
    # Chuyển đổi ma trận đầu vào từ chuỗi sang số
    input_matrix = data['matrix']
    comparison_matrix = []
    
    for row in input_matrix:
        numeric_row = []
        for value in row:
            if isinstance(value, str) and '/' in value:
                # Xử lý giá trị dạng phân số như "1/3", "1/5", etc.
                parts = value.split('/')
                if len(parts) == 2:
                    try:
                        numeric_value = float(parts[0]) / float(parts[1])
                        numeric_row.append(numeric_value)
                    except (ValueError, ZeroDivisionError):
                        return jsonify({"error": f"Không thể chuyển đổi giá trị '{value}' thành số"}), 400
            else:
                # Xử lý giá trị số hoặc chuỗi số
                try:
                    numeric_row.append(float(value))
                except ValueError:
                    return jsonify({"error": f"Không thể chuyển đổi giá trị '{value}' thành số"}), 400
        comparison_matrix.append(numeric_row)
    
    # Chuyển đổi sang numpy array và tính toán
    comparison_matrix = np.array(comparison_matrix)
    result = calculate_ahp(comparison_matrix)
    return jsonify(result)

@app.route('/calculate_alternative_matrices', methods=['POST'])
def calculate_alternative_matrices():
    """Tính toán ma trận đánh giá các phương án theo từng tiêu chí"""
    data = request.json
    criteria = data.get('criteria', [])
    alternatives = data.get('alternatives', [])
    matrices = data.get('matrices', {})
    criteria_weights = data.get('criteria_weights', [])
    
    # Kết quả của từng ma trận phương án theo từng tiêu chí
    alt_matrices_results = {}
    
    try:
        # Kiểm tra xem tất cả các tiêu chí đã có kết quả tính toán chưa
        for criterion in criteria:
            if criterion not in matrices:
                return jsonify({"error": f"Thiếu kết quả tính toán cho tiêu chí '{criterion}'"}), 400
            
            matrix_data = matrices[criterion]
            # Kiểm tra nếu matrix_data đã có thuộc tính weights và alternatives
            if isinstance(matrix_data, dict) and 'weights' in matrix_data and 'alternatives' in matrix_data and 'cr' in matrix_data:
                # Trường hợp đã có kết quả tính toán sẵn từ frontend
                alt_matrices_results[criterion] = matrix_data
            else:
                # Trường hợp là ma trận raw chưa tính toán (hiện tại không xảy ra do frontend đã tính)
                return jsonify({"error": f"Dữ liệu ma trận cho tiêu chí '{criterion}' không đúng định dạng"}), 400
    except Exception as e:
        return jsonify({"error": f"Lỗi khi xử lý ma trận: {str(e)}"}), 400
    
    # Tính điểm cuối cùng cho các phương án
    try:
        final_scores = calculate_final_scores(criteria_weights, alt_matrices_results)
        
        return jsonify({
            "alt_matrices_results": alt_matrices_results,
            "final_scores": final_scores
        })
    except Exception as e:
        return jsonify({"error": f"Lỗi khi tính điểm cuối cùng: {str(e)}"}), 400

@app.route('/get_alternative_matrices', methods=['POST'])
def get_alternative_matrices():
    """Lấy dữ liệu ma trận so sánh các phương án theo từng tiêu chí từ server"""
    try:
        data = request.json
        criteria = data.get('criteria', [])
        alternatives = data.get('alternatives', [])
        goal = data.get('goal', '')
        
        # Kiểm tra dữ liệu đầu vào
        if not criteria or not alternatives:
            return jsonify({"error": "Thiếu dữ liệu tiêu chí hoặc phương án"}), 400
            
        # Chuẩn bị ma trận so sánh cặp cho từng tiêu chí
        matrices = {}
        matrices_info = {}
        
        for criterion in criteria:
            n = len(alternatives)
            # Tạo ma trận với giá trị mặc định là 1 (đường chéo là 1, các giá trị khác được tính từ LLM)
            matrix = np.ones((n, n))
            
            # Tạo prompt để gửi tới LLM
            prompt = f"""
            Với mục tiêu: {goal}
            Hãy tạo ma trận so sánh cặp định lượng {n}x{n} cho các phương án sau đây theo tiêu chí "{criterion}":
            Các phương án: {', '.join(alternatives)}
            
            Sử dụng thang đo Saaty (1-9) với ý nghĩa:
            1: Hai phương án quan trọng như nhau
            3: Phương án A hơi quan trọng hơn B
            5: Phương án A quan trọng hơn B
            7: Phương án A rất quan trọng hơn B
            9: Phương án A cực kỳ quan trọng hơn B
            
            Hãy tính toán và cung cấp các giá trị dựa trên ý nghĩa định lượng của tiêu chí "{criterion}".
            Nếu tiêu chí mang tính định lượng (như giá, khoảng cách, thời gian, hiệu suất...), hãy sử dụng tỷ lệ giá trị thực để xác định độ quan trọng tương đối.
            Ví dụ:
            - Nếu tiêu chí là "Giá" và phương án A có giá 100, phương án B có giá 300, thì B kém hơn A khoảng 3 lần, do đó giá trị a_AB = 3 và a_BA = 1/3
            - Nếu tiêu chí là "Hiệu suất" và phương án A có hiệu suất 80%, phương án B có hiệu suất 40%, thì A tốt hơn B khoảng 2 lần, do đó giá trị a_AB = 2 và a_BA = 1/2
            
            Chỉ trả về ma trận dưới dạng JSON với định dạng:
            [
                [a_11, a_12, ..., a_1n],
                [a_21, a_22, ..., a_2n],
                ...
                [a_n1, a_n2, ..., a_nn]
            ]
            
            Hãy phân tích sâu các giá trị định lượng của từng phương án theo tiêu chí này và đánh giá chính xác nhất có thể.
            Nếu tiêu chí trừu tượng không có giá trị định lượng rõ ràng, hãy sử dụng thang đo Saaty và phân tích logic để đưa ra đánh giá hợp lý nhất.
            
            LƯU Ý QUAN TRỌNG: Ma trận phải đảm bảo tỷ số nhất quán CR < 0.1 (10%) theo phương pháp AHP.
            Đảm bảo tính nhất quán của ma trận bằng cách:
            1. Đường chéo ma trận phải có giá trị 1 (a_ii = 1)
            2. Ma trận phải đối xứng nghịch đảo: a_ji = 1/a_ij
            3. Tính chất bắc cầu: Nếu a_ij = x và a_jk = y, thì a_ik nên xấp xỉ x*y
            4. Hạn chế sử dụng các giá trị cực đoan (1/9 hoặc 9)
            """
            
            try:
                # Trước tiên thử gọi API để xác định các giá trị số học cho tiêu chí
                quantitative_prompt = f"""
                Phân tích tiêu chí "{criterion}" cho mục tiêu "{goal}":
                1. Tiêu chí này có phải là tiêu chí định lượng không? (có/không)
                2. Nếu là tiêu chí định lượng, hãy ước tính giá trị số cụ thể cho mỗi phương án: {', '.join(alternatives)}
                3. Đơn vị đo lường của tiêu chí này là gì?
                4. Tiêu chí này là tối đa hóa (càng cao càng tốt) hay tối thiểu hóa (càng thấp càng tốt)?
                
                Trả về kết quả dưới dạng JSON:
                {{
                    "is_quantitative": true/false,
                    "values": {{
                        "phương án 1": giá trị số,
                        "phương án 2": giá trị số,
                        ...
                    }},
                    "unit": "đơn vị",
                    "maximize": true/false
                }}
                
                Nếu không phải là tiêu chí định lượng, chỉ cần trả về:
                {{
                    "is_quantitative": false
                }}
                """
                
                response = requests.post(
                    API_URL,
                    json={
                        "messages": [
                            {"role": "system", "content": "Bạn là một chuyên gia phân tích dữ liệu. Hãy phân tích tính chất định lượng của tiêu chí được cung cấp."},
                            {"role": "user", "content": quantitative_prompt}
                        ],
                        "temperature": 0.3,
                        "max_tokens": 500
                    }
                )
                response.raise_for_status()
                
                is_quantitative = False
                quantitative_values = {}
                maximize = True  # Mặc định là tối đa hóa
                unit = ""
                
                response_data = response.json()
                if 'choices' in response_data and len(response_data['choices']) > 0:
                    content = response_data['choices'][0]['message']['content']
                    
                    # Trích xuất JSON từ kết quả
                    import json
                    import re
                    
                    # Tìm đối tượng JSON trong kết quả
                    json_match = re.search(r'\{.*\}', content, re.DOTALL)
                    if json_match:
                        try:
                            quant_data = json.loads(json_match.group(0))
                            is_quantitative = quant_data.get('is_quantitative', False)
                            
                            if is_quantitative:
                                quantitative_values = quant_data.get('values', {})
                                maximize = quant_data.get('maximize', True)
                                unit = quant_data.get('unit', "")
                                print(f"Phân tích định lượng cho tiêu chí '{criterion}':")
                                print(f"  Giá trị: {quantitative_values}")
                                print(f"  Tối đa hóa: {maximize}")
                                print(f"  Đơn vị: {unit}")
                        except json.JSONDecodeError as e:
                            print(f"Lỗi phân tích JSON quantitative: {e}")
                
                # Lưu thông tin định lượng
                if is_quantitative:
                    matrices_info[criterion] = {
                        "is_quantitative": True,
                        "maximize": maximize,
                        "unit": unit,
                        "values": quantitative_values
                    }
                else:
                    matrices_info[criterion] = {
                        "is_quantitative": False
                    }
                
                # Hàm để chuyển ma trận số thành ma trận string
                def convert_to_string_matrix(numeric_matrix):
                    string_matrix = []
                    for i in range(n):
                        string_row = []
                        for j in range(n):
                            val = numeric_matrix[i][j]
                            # Chuyển số thành string
                            if val == 1:
                                string_row.append("1")
                            elif val < 1:
                                # Chuyển số thập phân thành phân số 1/x
                                denominator = int(round(1/val))
                                string_row.append(f"1/{denominator}")
                            else:
                                string_row.append(str(int(round(val))))
                        string_matrix.append(string_row)
                    return string_matrix
                
                # Hàm để kiểm tra độ nhất quán của ma trận
                def check_consistency(matrix):
                    try:
                        result = calculate_ahp(np.array(matrix))
                        cr = result["cr"]
                        return 0 < cr < 0.1, cr
                    except Exception as e:
                        print(f"Lỗi kiểm tra độ nhất quán: {e}")
                        return False, 1.0
                
                # Hàm để điều chỉnh ma trận để có độ nhất quán tốt hơn
                def improve_consistency(matrix, orig_matrix=None):
                    # Nếu không có ma trận gốc, sử dụng ma trận hiện tại
                    if orig_matrix is None:
                        orig_matrix = matrix.copy()
                    
                    # Thử điều chỉnh các giá trị không trên đường chéo để cải thiện độ nhất quán
                    for _ in range(3):  # Thử tối đa 3 lần
                        # Tính độ nhất quán
                        result = calculate_ahp(matrix)
                        cr = result["cr"]
                        
                        # Nếu CR nằm trong khoảng mong muốn (0 < CR < 0.1), thì dừng
                        if 0 < cr < 0.1:
                            break
                            
                        # Nếu CR = 0, thêm một chút nhiễu để tạo sự khác biệt nhỏ
                        if cr == 0:
                            # Tạo ma trận nhiễu nhỏ
                            for i in range(n):
                                for j in range(i+1, n):  # Chỉ làm việc với nửa trên tam giác
                                    # Thêm nhiễu nhỏ, khoảng 5-10% của giá trị hiện tại
                                    noise_factor = 1.0 + (np.random.random() * 0.1 - 0.05)  # Nhiễu +/- 5%
                                    matrix[i][j] *= noise_factor
                                    matrix[j][i] = 1 / matrix[i][j]
                            
                            # Tính lại CR sau khi thêm nhiễu
                            result = calculate_ahp(matrix)
                            cr = result["cr"]
                            
                            # Nếu CR vẫn = 0, tạo một số bất đồng nhỏ ở ô giá trị lớn nhất
                            if cr == 0:
                                # Tìm giá trị lớn nhất trên tam giác trên
                                max_val = 0
                                max_i, max_j = 0, 0
                                for i in range(n):
                                    for j in range(i+1, n):
                                        if matrix[i][j] > max_val:
                                            max_val = matrix[i][j]
                                            max_i, max_j = i, j
                                
                                # Điều chỉnh giá trị lớn nhất nếu tìm thấy
                                if max_val > 1:
                                    matrix[max_i][max_j] = max_val * 1.15  # Tăng thêm 15%
                                    matrix[max_j][max_i] = 1 / matrix[max_i][max_j]
                        
                        # Tính eigenvalues và eigenvectors
                        weights = result["weights"]
                        
                        # Điều chỉnh ma trận dựa trên tính chất bắc cầu
                        for i in range(n):
                            for j in range(i+1, n):  # Chỉ làm việc với nửa trên tam giác
                                # Sử dụng tỷ lệ trọng số để điều chỉnh giá trị
                                ideal_value = weights[i] / weights[j]
                                
                                # Điều chỉnh giá trị phù hợp (giữ trong khoảng [1/9, 9])
                                ideal_value = min(9, max(1/9, ideal_value))
                                
                                # Điều chỉnh giá trị hiện tại về phía giá trị lý tưởng
                                current = matrix[i][j]
                                adjusted = current * 0.7 + ideal_value * 0.3  # Trộn 70% giá trị cũ, 30% giá trị tối ưu
                                
                                # Làm tròn về các giá trị Saaty (1,3,5,7,9 hoặc nghịch đảo)
                                if adjusted >= 1:
                                    saaty_vals = [1, 2, 3, 4, 5, 6, 7, 8, 9]
                                    closest_idx = min(range(len(saaty_vals)), key=lambda i: abs(saaty_vals[i] - adjusted))
                                    matrix[i][j] = saaty_vals[closest_idx]
                                    matrix[j][i] = 1 / matrix[i][j]
                                else:
                                    inv_adjusted = 1 / adjusted
                                    saaty_vals = [1, 2, 3, 4, 5, 6, 7, 8, 9]
                                    closest_idx = min(range(len(saaty_vals)), key=lambda i: abs(saaty_vals[i] - inv_adjusted))
                                    matrix[j][i] = saaty_vals[closest_idx]
                                    matrix[i][j] = 1 / matrix[j][i]
                    
                    # Kiểm tra lại CR sau khi điều chỉnh
                    result = calculate_ahp(matrix)
                    cr = result["cr"]
                    
                    # Nếu CR vẫn = 0, đảm bảo có sự khác biệt nhỏ giữa các giá trị
                    if cr == 0:
                        # Tạo một ma trận có CR > 0 nhưng vẫn nhỏ
                        for i in range(n):
                            for j in range(i+1, n):  # Chỉ làm việc với nửa trên tam giác
                                if i != j:
                                    # Thêm nhiễu khác nhau cho mỗi phần tử
                                    noise = 1.0 + (np.random.random() * 0.2 - 0.1)  # Nhiễu +/- 10%
                                    matrix[i][j] = max(1.0, matrix[i][j] * noise)
                                    if matrix[i][j] > 9: 
                                        matrix[i][j] = 9.0
                                    matrix[j][i] = 1.0 / matrix[i][j]
                    
                    return matrix
                
                # Nếu có dữ liệu định lượng, tạo ma trận dựa trên tỷ lệ giá trị
                if is_quantitative and quantitative_values and len(quantitative_values) >= len(alternatives):
                    # Tạo ma trận dựa trên tỷ lệ giá trị số
                    numeric_matrix = np.ones((n, n))
                    
                    for i, alt_i in enumerate(alternatives):
                        for j, alt_j in enumerate(alternatives):
                            if i == j:
                                # Đường chéo luôn là 1
                                numeric_matrix[i][j] = 1
                                continue
                                
                            if alt_i in quantitative_values and alt_j in quantitative_values:
                                val_i = float(quantitative_values[alt_i])
                                val_j = float(quantitative_values[alt_j])
                                
                                # Tránh chia cho 0
                                if val_i == 0 and val_j == 0:
                                    ratio = 1
                                elif val_i == 0:
                                    ratio = 1/9  # Giá trị thấp nhất nếu val_i = 0
                                elif val_j == 0:
                                    ratio = 9  # Giá trị cao nhất nếu val_j = 0
                                else:
                                    if maximize:
                                        ratio = val_i / val_j  # Càng cao càng tốt
                                    else:
                                        ratio = val_j / val_i  # Càng thấp càng tốt
                                
                                # Giới hạn tỷ lệ trong khoảng 1/9 đến 9 theo thang Saaty
                                if ratio > 9:
                                    ratio = 9
                                elif ratio < 1/9:
                                    ratio = 1/9
                                    
                                numeric_matrix[i][j] = ratio
                    
                    # Kiểm tra độ nhất quán
                    is_consistent, cr = check_consistency(numeric_matrix)
                    
                    # Nếu ma trận không nhất quán, thử điều chỉnh
                    if not is_consistent:
                        print(f"Ma trận ban đầu cho {criterion} không nhất quán (CR={cr}), đang điều chỉnh...")
                        numeric_matrix = improve_consistency(numeric_matrix)
                        is_consistent, cr = check_consistency(numeric_matrix)
                        print(f"Ma trận điều chỉnh cho {criterion}: CR={cr}")
                    
                    # Chuyển đổi về ma trận chuỗi
                    matrices[criterion] = convert_to_string_matrix(numeric_matrix)
                    matrices_info[criterion]["cr"] = cr
                    matrices_info[criterion]["is_consistent"] = is_consistent
                    
                    # Nếu đã tạo được ma trận nhất quán, tiếp tục với tiêu chí tiếp theo
                    if is_consistent:
                        continue
                
                # Nếu không có đủ dữ liệu định lượng hoặc ma trận vẫn không nhất quán,
                # tiếp tục với phương pháp LLM
                max_attempts = 3
                is_consistent = False
                cr = 1.0
                
                for attempt in range(max_attempts):
                    response = requests.post(
                        API_URL,
                        json={
                            "messages": [
                                {"role": "system", "content": "Bạn là một chuyên gia về phương pháp AHP. Hãy tạo ma trận so sánh cặp phương án theo tiêu chí được yêu cầu, đảm bảo CR < 0.1."},
                                {"role": "user", "content": prompt + f"\n\nĐây là lần thử thứ {attempt+1}. Tạo ma trận đảm bảo CR < 0.1."}
                            ],
                            "temperature": 0.3,
                            "max_tokens": 500
                        }
                    )
                    response.raise_for_status()
                    
                    response_data = response.json()
                    if 'choices' in response_data and len(response_data['choices']) > 0:
                        content = response_data['choices'][0]['message']['content']
                        
                        # Trích xuất ma trận từ kết quả
                        import json
                        import re
                        
                        # Tìm mảng JSON trong kết quả
                        matrix_match = re.search(r'\[\s*\[.*?\]\s*\]', content, re.DOTALL)
                        if matrix_match:
                            matrix_json = matrix_match.group(0)
                            try:
                                generated_matrix = json.loads(matrix_json)
                                
                                # Chuyển thành ma trận số để kiểm tra độ nhất quán
                                numeric_matrix = np.ones((n, n))
                                for i in range(n):
                                    for j in range(n):
                                        val = generated_matrix[i][j]
                                        if isinstance(val, str) and '/' in val:
                                            parts = val.split('/')
                                            if len(parts) == 2:
                                                try:
                                                    numeric_matrix[i][j] = float(parts[0]) / float(parts[1])
                                                except:
                                                    numeric_matrix[i][j] = 1
                                        else:
                                            try:
                                                numeric_matrix[i][j] = float(val)
                                            except:
                                                numeric_matrix[i][j] = 1
                                
                                # Kiểm tra độ nhất quán
                                is_consistent, cr = check_consistency(numeric_matrix)
                                
                                # Nếu ma trận không nhất quán, thử điều chỉnh
                                if not is_consistent:
                                    print(f"Ma trận LLM cho {criterion} lần {attempt+1} không nhất quán (CR={cr}), đang điều chỉnh...")
                                    numeric_matrix = improve_consistency(numeric_matrix)
                                    is_consistent, cr = check_consistency(numeric_matrix)
                                    print(f"Ma trận điều chỉnh cho {criterion}: CR={cr}")
                                
                                # Chuyển đổi về ma trận chuỗi
                                string_matrix = convert_to_string_matrix(numeric_matrix)
                                
                                matrices[criterion] = string_matrix
                                matrices_info[criterion]["cr"] = cr
                                matrices_info[criterion]["is_consistent"] = is_consistent
                                
                                if is_consistent:
                                    print(f"Tạo ma trận nhất quán cho {criterion} thành công sau {attempt+1} lần thử.")
                                    break
                            except Exception as e:
                                print(f"Lỗi xử lý ma trận cho {criterion}: {e}")
                
                # Nếu sau nhiều lần thử vẫn không có ma trận nhất quán, tạo ma trận đơn vị
                if not is_consistent:
                    print(f"Không thể tạo ma trận nhất quán cho {criterion}, sử dụng ma trận cơ bản.")
                    # Tạo ma trận đơn vị (đường chéo 1, các ô khác là giá trị nhỏ 2-3)
                    numeric_matrix = np.ones((n, n))
                    for i in range(n):
                        for j in range(i+1, n):
                            numeric_matrix[i][j] = 1  # Mặc định các ô không trên đường chéo là 1
                            numeric_matrix[j][i] = 1
                    
                    matrices[criterion] = convert_to_string_matrix(numeric_matrix)
                    matrices_info[criterion]["cr"] = 0.0
                    matrices_info[criterion]["is_consistent"] = True
                
            except Exception as e:
                print(f"Error generating matrix for {criterion}: {str(e)}")
                # Tạo ma trận mặc định nếu có lỗi
                numeric_matrix = np.ones((n, n))
                matrices[criterion] = convert_to_string_matrix(numeric_matrix)
                matrices_info[criterion]["cr"] = 0.0
                matrices_info[criterion]["is_consistent"] = True
        
        # Thêm trường thông tin về phân tích định lượng để hiển thị cho người dùng
        return jsonify({
            "matrices": matrices,
            "matrices_info": matrices_info,
            "matrix_info": "Ma trận so sánh được tạo dựa trên phân tích định lượng của các tiêu chí, đảm bảo tỷ số nhất quán CR < 10%."
        })
    
    except Exception as e:
        print(f"Error in get_alternative_matrices: {str(e)}")
        return jsonify({"error": f"Lỗi khi lấy ma trận: {str(e)}"}), 500

@app.route('/alternative_matrices')
def alternative_matrices():
    return render_template('alternative_matrices.html')

@app.route('/export-excel', methods=['POST'])
def export_excel():
    """Export AHP analysis results to Excel file with formulas"""
    try:
        data = request.json
        goal = data.get('goal', 'AHP Analysis')
        criteria = data.get('criteria', [])
        alternatives = data.get('alternatives', [])
        criteriaWeights = data.get('criteriaWeights', [])
        alternativeScores = data.get('alternativeScores', {})
        finalScores = data.get('finalScores', {})
        
        # Tạo tên file an toàn
        safe_filename = "".join([c for c in goal if c.isalpha() or c.isdigit() or c==' ']).rstrip()
        if len(safe_filename) > 30:
            safe_filename = safe_filename[:30]
        safe_filename = safe_filename.replace(' ', '_')
        
        # Create Excel writer
        output = io.BytesIO()
        writer = pd.ExcelWriter(output, engine='xlsxwriter')
        workbook = writer.book
        
        # Định nghĩa các formats
        header_format = workbook.add_format({
            'bold': True,
            'text_wrap': True,
            'valign': 'top',
            'fg_color': '#D7E4BC',
            'border': 1
        })
        
        percent_format = workbook.add_format({
            'num_format': '0.00%'
        })
        
        decimal_format = workbook.add_format({
            'num_format': '0.0000'
        })
        
        highlight_format = workbook.add_format({
            'bold': True,
            'fg_color': '#FFEB9C',
            'border': 1
        })
        
        # Tạo sheet "Input Data"
        input_df = pd.DataFrame({
            'Tiêu chí': criteria,
            'Trọng số': criteriaWeights
        })
        input_df.to_excel(writer, sheet_name='Input Data', index=False)
        input_worksheet = writer.sheets['Input Data']
        
        # Điều chỉnh độ rộng cột
        input_worksheet.set_column('A:A', 20)
        input_worksheet.set_column('B:B', 15)
        
        # Format header
        for col_num, value in enumerate(['Tiêu chí', 'Trọng số']):
            input_worksheet.write(0, col_num, value, header_format)
        
        # Format giá trị trọng số
        for row_num, weight in enumerate(criteriaWeights):
            input_worksheet.write(row_num + 1, 1, weight, decimal_format)
        
        # Tạo sheet "Ma trận tiêu chí"
        if alternativeScores:
            for i, criterion in enumerate(criteria):
                # Lấy dữ liệu cho tiêu chí hiện tại
                alt_weights = {}
                for alt in alternatives:
                    if alt in alternativeScores and criterion in alternativeScores[alt]:
                        alt_weights[alt] = alternativeScores[alt][criterion]
                    else:
                        alt_weights[alt] = 0
                
                # Tạo DataFrame và ghi vào sheet
                alt_df = pd.DataFrame({
                    'Phương án': alternatives,
                    f'Trọng số': [alt_weights[alt] for alt in alternatives]
                })
                alt_df.to_excel(writer, sheet_name=f'Criterion-{i+1}', index=False)
                alt_worksheet = writer.sheets[f'Criterion-{i+1}']
                
                # Thêm tên tiêu chí
                alt_worksheet.merge_range('A1:B1', f'Tiêu chí: {criterion}', header_format)
                alt_worksheet.write(1, 0, 'Phương án', header_format)
                alt_worksheet.write(1, 1, 'Trọng số', header_format)
                
                # Điều chỉnh độ rộng cột
                alt_worksheet.set_column('A:A', 20)
                alt_worksheet.set_column('B:B', 15)
                
                # Format giá trị
                for row_num, alt in enumerate(alternatives):
                    alt_worksheet.write(row_num + 2, 1, alt_weights[alt], decimal_format)
        
        # Tạo sheet "Calculations" với công thức
        calc_worksheet = workbook.add_worksheet('Calculations')
        
        # Viết tiêu đề
        calc_worksheet.merge_range('A1:E1', 'Tính toán kết quả AHP', header_format)
        calc_worksheet.write('A3', 'Tiêu chí', header_format)
        calc_worksheet.write('B3', 'Trọng số', header_format)
        
        # Viết tên phương án
        for col, alt in enumerate(alternatives):
            calc_worksheet.write(2, col + 2, alt, header_format)
        
        # Thêm dữ liệu tiêu chí và trọng số tiêu chí 
        for row, criterion in enumerate(criteria):
            # Tiêu chí
            calc_worksheet.write(row + 3, 0, criterion)
            # Trọng số tiêu chí - tham chiếu từ sheet Input Data
            cell_ref = f"='Input Data'!B{row + 2}"
            calc_worksheet.write_formula(row + 3, 1, cell_ref, decimal_format)
            
            # Trọng số phương án theo tiêu chí
            for col, alt in enumerate(alternatives):
                if alt in alternativeScores and criterion in alternativeScores[alt]:
                    # Tham chiếu từ sheet của tiêu chí tương ứng
                    cell_ref = f"='Criterion-{row+1}'!B{col + 3}"
                    calc_worksheet.write_formula(row + 3, col + 2, cell_ref, decimal_format)
                else:
                    calc_worksheet.write(row + 3, col + 2, 0, decimal_format)
        
        # Tính điểm tổng hợp với công thức
        row_sum = len(criteria) + 4
        calc_worksheet.merge_range(f'A{row_sum}:B{row_sum}', 'Điểm tổng hợp', header_format)
        
        # Viết công thức điểm tổng hợp cho từng phương án
        for col, alt in enumerate(alternatives):
            formula = f"=SUMPRODUCT(B4:B{len(criteria)+3},{xl_col_to_name(col+2)}4:{xl_col_to_name(col+2)}{len(criteria)+3})"
            calc_worksheet.write_formula(row_sum - 1, col + 2, formula, decimal_format)
        
        # Tính xếp hạng 
        row_rank = row_sum + 1
        calc_worksheet.merge_range(f'A{row_rank}:B{row_rank}', 'Xếp hạng', header_format)
        
        # Viết công thức xếp hạng
        rank_range = f"{xl_col_to_name(2)}{row_sum}:{xl_col_to_name(len(alternatives)+1)}{row_sum}"
        for col, alt in enumerate(alternatives):
            formula = f"=RANK({xl_col_to_name(col+2)}{row_sum},{rank_range},0)"
            calc_worksheet.write_formula(row_rank - 1, col + 2, formula)
        
        # Định dạng mở rộng
        calc_worksheet.set_column('A:A', 20)  # Cột tiêu chí
        calc_worksheet.set_column(f'B:{xl_col_to_name(len(alternatives)+1)}', 12)  # Các cột dữ liệu
        
        # Tạo sheet "Results"
        result_worksheet = workbook.add_worksheet('Results')
        
        # Viết tiêu đề
        result_worksheet.merge_range('A1:D1', f'Kết quả phân tích AHP: {goal}', header_format)
        result_worksheet.write('A3', 'Xếp hạng', header_format)
        result_worksheet.write('B3', 'Phương án', header_format)
        result_worksheet.write('C3', 'Điểm số', header_format)
        result_worksheet.write('D3', 'Tỷ lệ', header_format)
        
        # Viết công thức để lấy kết quả từ sheet Calculations
        for i, alt in enumerate(alternatives):
            row = i + 4
            # Xếp hạng
            result_worksheet.write_formula(f'A{row}', f"='Calculations'!{xl_col_to_name(i+2)}{row_rank}")
            # Phương án
            result_worksheet.write(f'B{row}', alt)
            # Điểm số
            result_worksheet.write_formula(f'C{row}', f"='Calculations'!{xl_col_to_name(i+2)}{row_sum}", decimal_format)
            # Tỷ lệ (so với tổng)
            result_worksheet.write_formula(f'D{row}', f"=C{row}/SUM('Calculations'!{xl_col_to_name(2)}{row_sum}:{xl_col_to_name(len(alternatives)+1)}{row_sum})", percent_format)
        
        # Sắp xếp kết quả theo xếp hạng
        result_worksheet.write('A2', 'Sắp xếp theo xếp hạng:')
        result_worksheet.data_validation('B2', {'validate': 'list', 'source': ['Tăng dần', 'Giảm dần']})
        result_worksheet.write('B2', 'Tăng dần')  # Giá trị mặc định
        
        # Thêm biểu đồ
        chart = workbook.add_chart({'type': 'bar'})
        chart.add_series({
            'name': 'Điểm số',
            'categories': ['Results', 3, 1, 3 + len(alternatives) - 1, 1],
            'values': ['Results', 3, 2, 3 + len(alternatives) - 1, 2],
            'data_labels': {'value': True},
        })
        chart.set_title({'name': f'Kết quả AHP: {goal}'})
        chart.set_x_axis({'name': 'Phương án'})
        chart.set_y_axis({'name': 'Điểm số'})
        result_worksheet.insert_chart('F3', chart, {'x_scale': 1.5, 'y_scale': 1.5})
        
        # Điều chỉnh độ rộng cột
        result_worksheet.set_column('A:A', 10)
        result_worksheet.set_column('B:B', 20)
        result_worksheet.set_column('C:D', 12)
        result_worksheet.set_column('E:E', 2)
        
        # Tạo sheet "Hướng dẫn"
        help_worksheet = workbook.add_worksheet('Hướng dẫn')
        help_worksheet.write('A1', 'Hướng dẫn sử dụng file Excel AHP', header_format)
        help_text = [
            'File Excel này chứa kết quả phân tích AHP với công thức động.',
            '',
            '1. Sheet "Input Data": Chứa trọng số các tiêu chí. Bạn có thể điều chỉnh các giá trị này.',
            '2. Sheet "Criterion-X": Chứa trọng số các phương án theo từng tiêu chí. Bạn cũng có thể điều chỉnh các giá trị này.',
            '3. Sheet "Calculations": Tính toán điểm tổng hợp sử dụng công thức SUMPRODUCT.',
            '4. Sheet "Results": Hiển thị kết quả xếp hạng các phương án với biểu đồ.',
            '',
            'Lưu ý:',
            '- Khi thay đổi giá trị ở sheet Input Data hoặc các sheet Criterion, kết quả sẽ tự động cập nhật.',
            '- Đảm bảo tổng các trọng số tiêu chí bằng 1 và tổng trọng số phương án theo từng tiêu chí cũng bằng 1.'
        ]
        for i, line in enumerate(help_text):
            help_worksheet.write(i+1, 0, line)
        help_worksheet.set_column('A:A', 100)
        
        # Đóng workbook và trả về file
        writer.close()
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f"AHP_Analysis_{safe_filename}.xlsx"
        )
        
    except Exception as e:
        print(f"Error exporting to Excel: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

def xl_col_to_name(col_num):
    """Chuyển đổi số cột Excel sang tên cột (0 -> A, 1 -> B, ...)"""
    result = ""
    while col_num >= 0:
        remainder = col_num % 26
        result = chr(65 + remainder) + result
        col_num = col_num // 26 - 1
    return result

@app.route('/export-pdf', methods=['POST'])
def export_pdf():
    """Export AHP analysis results to PDF file with professional, administrative report style"""
    try:
        data = request.json
        goal = data.get('goal', 'AHP Analysis')
        criteria = data.get('criteria', [])
        alternatives = data.get('alternatives', [])
        criteriaWeights = data.get('criteriaWeights', [])
        alternativeScores = data.get('alternativeScores', {})
        finalScores = data.get('finalScores', {})

        log_to_mongodb({
            'type': 'export_pdf',
            'goal': goal,
            'criteria': criteria,
            'alternatives': alternatives,
            'finalScores': finalScores
        })

        buffer = io.BytesIO()
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import os
        font_path = os.path.join(os.path.dirname(__file__), 'fonts', 'DejaVuSans.ttf')
        pdfmetrics.registerFont(TTFont('DejaVuSans', font_path))

        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
        from reportlab.platypus import Frame, PageTemplate

        doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
        styles = getSampleStyleSheet()
        elements = []

        # Custom styles
        styles.add(ParagraphStyle(
            name='Header', fontName='DejaVuSans', fontSize=14, alignment=TA_CENTER, textColor=colors.HexColor('#2563eb'), leading=18, spaceAfter=8, spaceBefore=0, fontWeight='bold', underlineWidth=1.5
        ))
        styles.add(ParagraphStyle(
            name='Title', fontName='DejaVuSans', fontSize=16, alignment=TA_CENTER, leading=22, spaceAfter=16, spaceBefore=0, fontWeight='bold'
        ))
        styles.add(ParagraphStyle(
            name='Section', fontName='DejaVuSans', fontSize=13, alignment=TA_CENTER, leading=18, spaceAfter=10, spaceBefore=10, textColor=colors.HexColor('#1e293b'), fontWeight='bold'
        ))
        styles.add(ParagraphStyle(
            name='TableHeader', fontName='DejaVuSans', fontSize=11, alignment=TA_CENTER, leading=14, textColor=colors.white, backColor=colors.HexColor('#2563eb'), spaceAfter=2, spaceBefore=2
        ))
        styles.add(ParagraphStyle(
            name='TableCell', fontName='DejaVuSans', fontSize=10, alignment=TA_LEFT, leading=13, spaceAfter=2, spaceBefore=2, wordWrap='CJK'
        ))
        styles.add(ParagraphStyle(
            name='TableCellCenter', fontName='DejaVuSans', fontSize=10, alignment=TA_CENTER, leading=13, spaceAfter=2, spaceBefore=2, wordWrap='CJK'
        ))
        styles.add(ParagraphStyle(
            name='Conclusion', fontName='DejaVuSans', fontSize=11, alignment=TA_JUSTIFY, leading=15, textColor=colors.HexColor('#166534'), backColor=colors.HexColor('#f0fdf4'), spaceAfter=10, spaceBefore=10, leftIndent=8, rightIndent=8, borderPadding=6
        ))
        styles.add(ParagraphStyle(
            name='Footer', fontName='DejaVuSans', fontSize=9, alignment=TA_CENTER, textColor=colors.HexColor('#64748b'), leading=12, spaceBefore=12
        ))

        # Header
        elements.append(Paragraph("NHÓM 7 - HỆ HỖ TRỢ RA QUYẾT ĐỊNH", styles['Header']))
        elements.append(Spacer(1, 0.08*inch))
        # Title
        elements.append(Paragraph(f"Phân tích AHP: {goal}", styles['Title']))
        elements.append(Spacer(1, 0.12*inch))

        # Section: Criteria Weights
        elements.append(Paragraph("Trọng số các tiêu chí", styles['Section']))
        # Table data
        criteria_data = [
            [Paragraph('Tiêu chí', styles['TableHeader']),
             Paragraph('Trọng số', styles['TableHeader']),
             Paragraph('Tỷ lệ', styles['TableHeader'])]
        ]
        for i, criterion in enumerate(criteria):
            criteria_data.append([
                Paragraph(criterion, styles['TableCell']),
                Paragraph(f"{criteriaWeights[i]:.4f}", styles['TableCellCenter']),
                Paragraph(f"{criteriaWeights[i]*100:.2f}%", styles['TableCellCenter'])
            ])
        criteria_table = Table(criteria_data, colWidths=[2.8*inch, 1.1*inch, 1.1*inch], repeatRows=1)
        criteria_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), 'DejaVuSans'),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2563eb')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('ALIGN', (1, 1), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.whitesmoke, colors.HexColor('#e0e7ff')]),
            ('GRID', (0, 0), (-1, -1), 0.7, colors.HexColor('#64748b')),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        elements.append(criteria_table)
        elements.append(Spacer(1, 0.18*inch))

        # Section: Final Results
        elements.append(Paragraph("Kết quả phân tích cuối cùng", styles['Section']))
        # Sort alternatives by final score
        sorted_alternatives = sorted(alternatives, key=lambda alt: finalScores.get(alt, 0), reverse=True)
        total_score = sum(finalScores.values())
        final_data = [
            [Paragraph('Xếp hạng', styles['TableHeader']),
             Paragraph('Phương án', styles['TableHeader']),
             Paragraph('Điểm số', styles['TableHeader']),
             Paragraph('Tỷ lệ', styles['TableHeader'])]
        ]
        for i, alt in enumerate(sorted_alternatives):
            score = finalScores.get(alt, 0)
            percentage = (score / total_score) * 100 if total_score > 0 else 0
            row = [
                Paragraph(f"{i+1}", styles['TableCellCenter']),
                Paragraph(alt, styles['TableCell']),
                Paragraph(f"{score:.4f}", styles['TableCellCenter']),
                Paragraph(f"{percentage:.2f}%", styles['TableCellCenter'])
            ]
            final_data.append(row)
        final_table = Table(final_data, colWidths=[0.8*inch, 3.2*inch, 1.1*inch, 1.1*inch], repeatRows=1)
        # Highlight best row
        best_row_color = colors.HexColor('#bbf7d0')
        final_table_style = [
            ('FONTNAME', (0, 0), (-1, -1), 'DejaVuSans'),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2563eb')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.whitesmoke, colors.HexColor('#e0e7ff')]),
            ('GRID', (0, 0), (-1, -1), 0.7, colors.HexColor('#64748b')),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('ALIGN', (0, 1), (0, -1), 'CENTER'),
            ('ALIGN', (2, 1), (3, -1), 'CENTER'),
            ('ALIGN', (1, 1), (1, -1), 'LEFT'),
            ('BACKGROUND', (0, 1), (-1, 1), best_row_color),
        ]
        final_table.setStyle(TableStyle(final_table_style))
        elements.append(final_table)
        elements.append(Spacer(1, 0.18*inch))

        # Conclusion
        if sorted_alternatives:
            best_alt = sorted_alternatives[0]
            best_score = finalScores.get(best_alt, 0)
            percentage = (best_score / total_score) * 100 if total_score > 0 else 0
            conclusion_text = f"<b>Kết luận:</b> Phương án <b>{best_alt}</b> là lựa chọn tốt nhất với điểm số <b>{best_score:.4f}</b> (<b>{percentage:.2f}%</b>)."
            if len(sorted_alternatives) > 1:
                second_alt = sorted_alternatives[1]
                second_score = finalScores.get(second_alt, 0)
                diff_percent = ((best_score - second_score) / best_score) * 100 if best_score > 0 else 0
                conclusion_text += f" Phương án này vượt trội hơn phương án <b>{second_alt}</b> khoảng <b>{diff_percent:.2f}%</b>."
            elements.append(Spacer(1, 0.05*inch))
            elements.append(Paragraph(conclusion_text, styles['Conclusion']))

        # Footer
        elements.append(Spacer(1, 0.18*inch))
        footer_text = f"Nhóm 7 - Hệ hỗ trợ ra quyết định | Báo cáo được tạo vào: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}"
        elements.append(Paragraph(footer_text, styles['Footer']))

        doc.build(elements)
        buffer.seek(0)
        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f"AHP_Analysis_{goal.replace(' ', '_')[:30]}.pdf"
        )
    except Exception as e:
        print(f"Error exporting to PDF: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/save-log', methods=['POST'])
def save_log():
    try:
        data = request.json
        log_to_mongodb({
            'type': 'manual_save',
            **data
        })
        return jsonify({"success": True, "message": "Đã lưu kết quả vào MongoDB!"})
    except Exception as e:
        return jsonify({"success": False, "message": f"Lỗi: {str(e)}"}), 500

app.run(debug=True) 